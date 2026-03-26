"""
DOCX Report Generator — combines base template + report data into final document.

Pipeline:
1. Load base template (with Content Controls)
2. Update SDT values with user placeholders
3. Replace {{REPORT_SECTIONS}} marker with rendered sections
4. Save final document

Uses python-docx + lxml for SDT manipulation.
Markdown-to-DOCX rendering adapted from existing report_generator.py patterns.
"""
from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

from backend.report_template_generator import (
    _find_and_update_sdt,
    generate_gsm_template,
    generate_aml_template,
)


# ─── Inline markdown regex (from report_generator.py) ────────────────────────

_MD_INLINE_RE = re.compile(
    r'(\*\*\*(.+?)\*\*\*'        # ***bold italic***
    r'|\*\*(.+?)\*\*'             # **bold**
    r'|\*(.+?)\*'                 # *italic*
    r'|`(.+?)`'                   # `code`
    r')',
)


def _add_md_runs(para, text: str) -> None:
    """Add runs to paragraph honouring inline markdown (bold, italic, code)."""
    last = 0
    for m in _MD_INLINE_RE.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(2):          # ***bold+italic***
            r = para.add_run(m.group(2))
            r.bold = True
            r.italic = True
        elif m.group(3):        # **bold**
            r = para.add_run(m.group(3))
            r.bold = True
        elif m.group(4):        # *italic*
            r = para.add_run(m.group(4))
            r.italic = True
        elif m.group(5):        # `code`
            r = para.add_run(m.group(5))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


# ─── Markdown table parser (from report_generator.py) ────────────────────────

def _parse_md_table(lines: List[str]) -> List[List[str]]:
    """Parse markdown table lines into list of rows (list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        # Skip separator rows (---|---|---)
        if all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
            continue
        rows.append(cells)
    return rows


# ─── Main generator ──────────────────────────────────────────────────────────

def generate_docx_report(
    template_path: Optional[Path],
    report_data: dict,
    placeholders: Dict[str, str],
    output_path: Path,
    report_type: str = "gsm",
) -> Path:
    """Generate a DOCX report from template + data.

    Args:
        template_path: Path to base .docx template. If None, generates default.
        report_data: Output from report_builder.build_report_data().
        placeholders: User-filled SDT values (INSTYTUCJA, SYGNATURA, etc.)
        output_path: Where to save the final .docx file.
        report_type: "gsm" or "aml"

    Returns:
        Path to saved DOCX file.
    """
    # ── Step 1: Load or generate template ──
    if template_path and template_path.exists():
        doc = Document(str(template_path))
    else:
        # Generate default template in-memory
        if report_type == "aml":
            template_bytes = generate_aml_template(placeholders)
        else:
            template_bytes = generate_gsm_template(placeholders)
        doc = Document(BytesIO(template_bytes))

    body = doc.element.body

    # ── Step 2: Update Content Controls with placeholder values ──
    for tag_name, value in placeholders.items():
        if value:  # Only update non-empty values
            _find_and_update_sdt(body, tag_name, value)

    # ── Step 3: Find and replace {{REPORT_SECTIONS}} marker ──
    marker_found = False
    marker_paragraph = None

    for p_elem in body.findall(qn("w:p")):
        # Check all text in this paragraph
        full_text = ""
        for r_elem in p_elem.findall(qn("w:r")):
            for t_elem in r_elem.findall(qn("w:t")):
                if t_elem.text:
                    full_text += t_elem.text

        if "{{REPORT_SECTIONS}}" in full_text:
            marker_paragraph = p_elem
            marker_found = True
            break

    if marker_found and marker_paragraph is not None:
        # Get the index of the marker paragraph
        parent = marker_paragraph.getparent()
        marker_idx = list(parent).index(marker_paragraph)

        # Remove the marker paragraph
        parent.remove(marker_paragraph)

        # Insert sections at marker position
        insert_idx = marker_idx
        for section in report_data.get("sections", []):
            elements = _render_section_to_docx_elements(doc, section)
            for elem in elements:
                parent.insert(insert_idx, elem)
                insert_idx += 1
    else:
        # No marker found — append sections at the end (before footer SDTs)
        # Find last separator line to insert before it
        for section in report_data.get("sections", []):
            elements = _render_section_to_docx_elements(doc, section)
            for elem in elements:
                # Insert before the last few elements (footer)
                body.append(elem)

    # ── Step 4: Save ──
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ─── Section-to-DOCX rendering ───────────────────────────────────────────────

def _render_section_to_docx_elements(doc: Document, section: dict) -> List:
    """Render a single report section to a list of DOCX XML elements.

    Uses content_md (markdown) as the primary source. Falls back to tables.
    """
    elements = []
    title = section.get("title", "Sekcja")
    content_md = section.get("content_md", "")
    tables = section.get("tables", [])

    # Section heading
    heading = doc.add_heading(title, level=2)
    elements.append(heading._element)

    # Render markdown content
    if content_md:
        md_elements = _render_markdown_to_elements(doc, content_md)
        elements.extend(md_elements)

    # Render structured tables (if any, and not already in markdown)
    if tables and not content_md:
        for tbl_def in tables:
            tbl_elements = _render_table_to_elements(doc, tbl_def)
            elements.extend(tbl_elements)

    # Spacing after section
    spacer = doc.add_paragraph("")
    elements.append(spacer._element)

    return elements


def _render_markdown_to_elements(doc: Document, md_text: str) -> List:
    """Parse markdown text and create DOCX paragraph/table elements."""
    elements = []
    lines = md_text.replace("\r\n", "\n").split("\n")

    in_code = False
    code_buf: List[str] = []
    table_buf: List[str] = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph("\n".join(code_buf))
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elements.append(p._element)
        code_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        parsed = _parse_md_table(table_buf)
        table_buf = []
        if not parsed:
            return
        cols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=cols)
        tbl.style = "Table Grid"
        for r_i, row in enumerate(parsed):
            for c_i in range(cols):
                cell_text = row[c_i] if c_i < len(row) else ""
                tbl.cell(r_i, c_i).text = cell_text
                # Bold header row
                if r_i == 0:
                    for run in tbl.cell(r_i, c_i).paragraphs[0].runs:
                        run.bold = True
        elements.append(tbl._element)

    for line in lines:
        stripped = line.strip()

        # Fenced code blocks
        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        # Table rows
        if stripped.startswith("|"):
            table_buf.append(stripped)
            continue
        elif table_buf:
            flush_table()

        # Empty lines
        if not stripped:
            continue

        # Headings (sub-section)
        if stripped.startswith("####"):
            p = doc.add_heading(stripped.lstrip("#").strip(), level=4)
            elements.append(p._element)
            continue
        if stripped.startswith("###"):
            p = doc.add_heading(stripped.lstrip("#").strip(), level=3)
            elements.append(p._element)
            continue

        # Bullet list
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_md_runs(p, stripped[2:])
            elements.append(p._element)
            continue

        # Numbered list
        m_num = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _add_md_runs(p, m_num.group(2))
            elements.append(p._element)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_md_runs(p, stripped)
        elements.append(p._element)

    # Flush any remaining
    flush_code()
    flush_table()

    return elements


def _render_table_to_elements(doc: Document, tbl_def: dict) -> List:
    """Render a structured table definition to DOCX table element."""
    elements = []
    headers = tbl_def.get("headers", [])
    rows = tbl_def.get("rows", [])
    title = tbl_def.get("title", "")

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        elements.append(p._element)

    if not rows:
        return elements

    cols = len(headers) if headers else (len(rows[0]) if rows else 0)
    total_rows = (1 if headers else 0) + len(rows)

    tbl = doc.add_table(rows=total_rows, cols=cols)
    tbl.style = "Table Grid"

    row_offset = 0
    if headers:
        for c_i, h in enumerate(headers):
            tbl.cell(0, c_i).text = str(h)
            for run in tbl.cell(0, c_i).paragraphs[0].runs:
                run.bold = True
        row_offset = 1

    for r_i, row in enumerate(rows):
        for c_i in range(cols):
            val = row[c_i] if c_i < len(row) else ""
            tbl.cell(r_i + row_offset, c_i).text = str(val)

    elements.append(tbl._element)
    return elements
