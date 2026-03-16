"""
GSM Note DOCX Generator — fills the professional note template with data.

Pipeline:
  1. docxtpl renders Jinja2 placeholders
  2. python-docx post-processes:
     - removes old anomaly synthetic description paragraphs
     - inserts data tables after section headings
     - inserts descriptive paragraphs (anomaly intro, contact graph text, etc.)
     - inserts chart images with captions after relevant sections
     - inserts bullet-list items for locations and movement data
     - applies consistent formatting (justified, same font, orphan prevention)
     - adds page numbers in footer (page/total, right-aligned)
"""
from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

log = logging.getLogger("aistate.gsm.note_gen")

# ─── Constants ───────────────────────────────────────────────────────────────

FONT_NAME = "Calibri"
FONT_SIZE_BODY = 10   # pt — body text
FONT_SIZE_TABLE = 8   # pt — table cells
FONT_SIZE_TABLE_SMALL = 7  # pt — small table cells (anomaly Dane)
FONT_SIZE_CAPTION = 9  # pt — table/chart captions

# Polish single-letter words that shouldn't be left at line end
_ORPHAN_RE = re.compile(r'\b([AaIiOoUuWwZz]) ', re.UNICODE)


# ─── Table definitions ──────────────────────────────────────────────────────

TABLE_DEFS = {
    "stats": {
        "section_heading": "4.",
        "caption": "Tabela: Statystyki aktywno\u015bci telekomunikacyjnej",
        "columns": ["Parametr", "Warto\u015b\u0107"],
    },
    "contacts": {
        "section_heading": "5.",
        "caption": "Tabela: Najcz\u0119stsze kontakty",
        "columns": ["Numer", "Interakcje", "Po\u0142. wych.", "Po\u0142. przych."],
    },
    "anomalies": {
        "section_heading": "7.",  # was 6, now 7 after "Rodzaj aktywności" inserted as 6
        "caption": "Tabela: Wykryte anomalie",
        "columns": ["Kategoria", "Opis", "Dane"],
    },
    "locations": {
        "section_heading": "8.",  # was 7, now 8
        "caption": "Tabela: Lokalizacje BTS",
        "columns": ["Lokalizacja", "Liczba rekord\u00f3w", "Udzia\u0142 %"],
    },
}

# ─── Descriptive text blocks ────────────────────────────────────────────────

ANOMALY_INTRO_TEXT = (
    "W\u00a0toku analizy materia\u0142u bilingowego zidentyfikowano zdarzenia oraz "
    "sekwencje aktywno\u015bci odbiegaj\u0105ce od typowego profilu korzystania "
    "z\u00a0numeru. Za anomalie uznano zar\u00f3wno pojedyncze incydenty "
    "o\u00a0niestandardowych cechach, jak i\u00a0powtarzalne wzorce czasowe, "
    "kontaktowe lub lokalizacyjne, kt\u00f3re mog\u0105 wymaga\u0107 "
    "pog\u0142\u0119bionej weryfikacji analitycznej."
)

CONTACTS_GRAPH_TEXT = (
    "Graf \u201eNajcz\u0119stsze kontakty\u201d obrazuje struktur\u0119 relacji "
    "komunikacyjnych analizowanego numeru, wskazuj\u0105c podmioty wyst\u0119puj\u0105ce "
    "najcz\u0119\u015bciej w\u00a0ruchu telekomunikacyjnym. Pozwala to okre\u015bli\u0107 "
    "kr\u0105g najaktywniejszych kontakt\u00f3w, uchwyci\u0107 powtarzalno\u015b\u0107 "
    "komunikacji oraz wskaza\u0107 numery, kt\u00f3re mog\u0105 odgrywa\u0107 istotn\u0105 "
    "rol\u0119 w\u00a0badanym modelu \u0142\u0105czno\u015bci."
)

ACTIVITY_DISTRIBUTION_TEXT = (
    "Mapa rozk\u0142adu aktywno\u015bci przedstawia intensywno\u015b\u0107 zdarze\u0144 "
    "telekomunikacyjnych w\u00a0zale\u017cno\u015bci od dnia tygodnia i\u00a0pory doby. "
    "Zestawienie to pozwala uchwyci\u0107 dominuj\u0105ce przedzia\u0142y aktywno\u015bci, "
    "wskaza\u0107 powtarzalne schematy czasowe oraz oceni\u0107, czy komunikacja "
    "koncentrowa\u0142a si\u0119 w\u00a0typowych godzinach dziennych, czy r\u00f3wnie\u017c "
    "w\u00a0porach nietypowych."
)

# Chart insertion order
CHART_INSERTION_ORDER = [
    {
        "key": "top_contacts",
        "caption": "Graf: Najcz\u0119stsze kontakty",
        "pre_text": CONTACTS_GRAPH_TEXT,
        "section_heading": "5.",
    },
    {
        "key": "activity",
        "caption": "Rozk\u0142ad aktywno\u015bci",
        "pre_text": ACTIVITY_DISTRIBUTION_TEXT,
        "new_section_heading": "6. Rodzaj aktywno\u015bci",
        "section_heading": "5.",
    },
    {
        "key": "night_activity",
        "caption": "Aktywno\u015b\u0107 nocna",
        "section_heading": "5.",  # inserted after contacts, before anomalies
    },
    {
        "key": "weekend_activity",
        "caption": "Aktywno\u015b\u0107 weekendowa",
        "section_heading": "5.",
    },
    {
        "key": "map_bts",
        "caption": "Mapa lokalizacji BTS",
        "section_heading": "8.",  # was 7, now 8 after renumbering
    },
]

# Texts to remove from the rendered DOCX (old anomaly synthetic description)
_REMOVE_TEXT_MARKERS = [
    "W raporcie wygenerowano",
    "Ich syntetyczny opis obejmuje",
    "Na szczeg\u00f3ln\u0105 uwag\u0119 zas\u0142uguj\u0105",
    "Ujawnione wzorce mog\u0105 wskazywa\u0107",
]


# ─── Public API ─────────────────────────────────────────────────────────────

def generate_note_docx(
    template_path: Path,
    placeholders: Dict[str, Any],
    output_path: Path,
    *,
    chart_images: Optional[Dict[str, bytes]] = None,
    llm_overrides: Optional[Dict[str, str]] = None,
    table_data: Optional[Dict[str, List[List[str]]]] = None,
    selected_tables: Optional[List[str]] = None,
    custom_template: Optional[Dict[str, Any]] = None,
) -> Path:
    """Generate a professional analytical note DOCX from the template."""
    from docxtpl import DocxTemplate

    if llm_overrides:
        _apply_llm_overrides(placeholders, llm_overrides)

    # Extract internal data (prefixed with _) before flattening
    anomaly_table_rows = placeholders.pop("_anomaly_table_rows", [])
    location_areas_list = placeholders.pop("_location_areas_list", [])
    location_movement_list = placeholders.pop("_location_movement_list", [])
    contacts_bullet_list = placeholders.pop("_contacts_bullet_list", [])
    special_numbers_footnotes = placeholders.pop("_special_numbers_footnotes", [])

    tpl = DocxTemplate(str(template_path))
    context = _flatten_for_template(placeholders)
    tpl.render(context)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(output_path))

    # Extract custom text overrides from user template
    tpl_texts = _extract_template_texts(custom_template) if custom_template else {}

    # Post-processing with python-docx
    _post_process_docx(
        output_path,
        table_data=table_data,
        selected_tables=selected_tables,
        anomaly_table_rows=anomaly_table_rows,
        chart_images=chart_images,
        location_areas_list=location_areas_list,
        location_movement_list=location_movement_list,
        contacts_bullet_list=contacts_bullet_list,
        special_numbers_footnotes=special_numbers_footnotes,
        tpl_texts=tpl_texts,
    )

    log.info("Generated GSM note: %s (%d bytes)", output_path.name, output_path.stat().st_size)
    return output_path


# ─── Post-processing ────────────────────────────────────────────────────────

def _post_process_docx(
    docx_path: Path,
    *,
    table_data: Optional[Dict[str, List[List[str]]]] = None,
    selected_tables: Optional[List[str]] = None,
    anomaly_table_rows: Optional[List[List[str]]] = None,
    chart_images: Optional[Dict[str, bytes]] = None,
    location_areas_list: Optional[List[str]] = None,
    location_movement_list: Optional[List[str]] = None,
    contacts_bullet_list: Optional[List[str]] = None,
    special_numbers_footnotes: Optional[List[str]] = None,
    tpl_texts: Optional[Dict[str, str]] = None,
) -> None:
    """Insert tables, charts, texts, lists; fix formatting; add page numbers."""
    from docx import Document

    doc = Document(str(docx_path))
    body = doc.element.body

    # 0a. Renumber sections 6→7, 7→8, 8→9, 9→10, 10→11 (to make room for "6. Rodzaj aktywności")
    _renumber_sections(body)

    # 0b. Remove old anomaly synthetic description paragraphs
    _remove_old_anomaly_text(body)

    # 0b. Fix "za okres od" — move to new line
    _fix_za_okres_newline(body)

    # 0c. Replace "Data raportu HTML" with "Data sporz\u0105dzenia"
    _replace_text_in_tables(doc, "Data raportu HTML", "Data sporz\u0105dzenia")

    # 1. Insert anomaly intro text + table with Kategoria/Opis/Dane
    if anomaly_table_rows:
        anomaly_intro = (tpl_texts or {}).get("anomaly_intro", ANOMALY_INTRO_TEXT)
        _insert_anomaly_section(doc, body, anomaly_table_rows, intro_text=anomaly_intro)

    # 2. Insert standard tables (stats, contacts, locations)
    if selected_tables and table_data:
        tables_to_insert = [t for t in selected_tables if t != "anomalies"]
        if tables_to_insert:
            _insert_tables(doc, body, table_data, tables_to_insert)

    # 3. Insert location bullet-lists (all items end with ";")
    if location_areas_list:
        _insert_bullet_list_after_text(
            doc, body,
            search_text="koncentrowa\u0142a si\u0119 w rejonach",
            items=location_areas_list,
        )
    if location_movement_list:
        _insert_bullet_list_after_text(
            doc, body,
            search_text="przes\u0142anki dotycz\u0105ce przemieszczania",
            items=location_movement_list,
        )

    # 3b. Insert contacts bullet-list after "Najcz\u0119\u015bciej wyst\u0119puj\u0105ce kontakty"
    if contacts_bullet_list:
        _insert_bullet_list_after_text(
            doc, body,
            search_text="Najcz\u0119\u015bciej wyst\u0119puj\u0105ce kontakty",
            items=contacts_bullet_list,
        )

    # 3c. Insert special numbers footnotes
    if special_numbers_footnotes:
        _insert_footnotes_after_text(
            doc, body,
            search_text="numer\u00f3w specjalnych odnotowano",
            footnotes=special_numbers_footnotes,
        )

    # 3d. Insert BTS count + dominant activity sentence after location lists
    _insert_bts_summary_sentence(doc, body, location_areas_list)

    # 4. Insert chart images
    if chart_images:
        _insert_chart_images(doc, body, chart_images, tpl_texts=tpl_texts)

    # 5. Fix formatting: consistent font, justified text, orphan prevention
    _fix_formatting(doc)

    # 6. Add page numbers in footer
    _add_page_numbers(doc)

    doc.save(str(docx_path))


# ─── Remove old anomaly text ────────────────────────────────────────────────

def _renumber_sections(body) -> None:
    """Renumber sections 6→7, 7→8, 8→9, 9→10, 10→11 in headings.

    This makes room for the new '6. Rodzaj aktywności' section.
    Must be called before other post-processing that looks for section numbers.
    """
    from docx.oxml.ns import qn

    # Process in reverse order so we don't double-renumber (10→11 first, then 9→10, etc.)
    renumber_map = [
        ("10.", "11."),
        ("9.", "10."),
        ("8.", "9."),
        ("7.", "8."),
        ("6.", "7."),
    ]

    for child in body:
        if child.tag == qn('w:p'):
            pStyle = child.find(f'.//{qn("w:pStyle")}')
            is_heading = pStyle is not None and 'Heading' in pStyle.get(qn('w:val'), '')
            if not is_heading:
                continue

            full_text = ''.join(child.itertext()).strip()
            for old_prefix, new_prefix in renumber_map:
                if full_text.startswith(old_prefix):
                    # Find the text element and replace prefix
                    for t_el in child.iter(qn('w:t')):
                        if t_el.text and t_el.text.strip().startswith(old_prefix):
                            t_el.text = t_el.text.replace(old_prefix, new_prefix, 1)
                            log.debug("Renumbered section '%s' → '%s'", old_prefix, new_prefix)
                            break
                    break


def _remove_old_anomaly_text(body) -> None:
    """Remove paragraphs containing the old anomaly synthetic description."""
    from docx.oxml.ns import qn

    to_remove = []
    for child in body:
        if child.tag == qn('w:p'):
            full_text = ''.join(child.itertext()).strip()
            if not full_text:
                continue
            for marker in _REMOVE_TEXT_MARKERS:
                if marker in full_text:
                    to_remove.append(child)
                    break

    for elem in to_remove:
        body.remove(elem)
        log.debug("Removed old anomaly text paragraph")


def _fix_za_okres_newline(body) -> None:
    """Move 'za okres od' to a new line by inserting a line break before it."""
    from docx.oxml.ns import qn
    from lxml import etree

    for child in body:
        if child.tag == qn('w:p'):
            for run_el in child.findall(f'.//{qn("w:r")}'):
                for t_el in run_el.findall(qn('w:t')):
                    if t_el.text and 'za okres od' in t_el.text:
                        # Split text at "za okres od" and insert line break
                        parts = t_el.text.split('za okres od', 1)
                        t_el.text = parts[0].rstrip()
                        # Create break element
                        br = etree.SubElement(run_el, qn('w:br'))
                        # Create new text element after break
                        new_t = etree.SubElement(run_el, qn('w:t'))
                        new_t.text = 'za okres od' + parts[1]
                        new_t.set(qn('xml:space'), 'preserve')
                        log.debug("Inserted line break before 'za okres od'")
                        return


def _replace_text_in_tables(doc, old_text: str, new_text: str) -> None:
    """Replace text in all table cells."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            log.debug("Replaced '%s' with '%s'", old_text, new_text)


def _insert_footnotes_after_text(doc, body, search_text: str, footnotes: List[str]) -> None:
    """Insert footnote-style explanations as small text after a paragraph."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    target = None
    for child in body:
        if child.tag == qn('w:p'):
            full_text = ''.join(child.itertext()).strip()
            if search_text in full_text:
                target = child
                break

    if target is None:
        log.debug("Text '%s' not found for footnotes", search_text[:40])
        return

    # Insert footnote paragraph
    fn_p = _insert_paragraph_after(doc, body, target)
    fn_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fn_p.paragraph_format.left_indent = Cm(0.5)
    for i, fn_text in enumerate(footnotes):
        if i > 0:
            fn_p.add_run("\n")
        superscript_run = fn_p.add_run(f"{i + 1}) ")
        _style_run(superscript_run, 7, italic=True, color=(80, 80, 80))
        text_run = fn_p.add_run(_fix_orphans(fn_text))
        _style_run(text_run, 7, italic=True, color=(80, 80, 80))

    log.debug("Inserted %d footnotes after '%s'", len(footnotes), search_text[:40])


def _insert_bts_summary_sentence(doc, body, location_areas_list: Optional[List[str]]) -> None:
    """Insert 'Łączna liczba unikalnych lokalizacji BTS...' after location bullet-list."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    if not location_areas_list:
        return

    bts_count = len(location_areas_list)
    dominant = location_areas_list[0] if location_areas_list else "brak danych"
    # Strip record count from dominant location name
    if " (" in dominant:
        dominant = dominant.split(" (")[0]

    # Find the last bullet-list item inserted after "koncentrowała się w rejonach"
    target = None
    for child in body:
        if child.tag == qn('w:p'):
            full_text = ''.join(child.itertext()).strip()
            if full_text.startswith("\u2022") and full_text.endswith(";"):
                target = child  # keep updating — last bullet wins

    if target is None:
        log.debug("No bullet-list items found for BTS summary")
        return

    p = _insert_paragraph_after(doc, body, target)
    text = _fix_orphans(
        f"\u0141\u0105czna liczba unikalnych lokalizacji BTS wynios\u0142a {bts_count}, "
        f"za\u015b dominuj\u0105ce punkty aktywno\u015bci opisano jako: {dominant}."
    )
    run = p.add_run(text)
    _style_run(run, FONT_SIZE_BODY)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    log.debug("Inserted BTS summary sentence")


# ─── Anomaly section ────────────────────────────────────────────────────────

def _insert_anomaly_section(doc, body, anomaly_rows: List[List[str]], *, intro_text: Optional[str] = None) -> None:
    """Insert anomaly intro text and Kategoria/Opis/Dane table in section 6."""
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    insert_after = _find_section_heading(body, "7.")
    if insert_after is None:
        log.warning("Section 7 heading not found, skipping anomaly section")
        return

    # Intro text
    intro_p = _insert_paragraph_after(doc, body, insert_after)
    text = intro_text if intro_text else ANOMALY_INTRO_TEXT
    run = intro_p.add_run(_fix_orphans(text))
    _style_run(run, FONT_SIZE_BODY)
    intro_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Caption
    caption_p = _insert_paragraph_after(doc, body, intro_p._element)
    run = caption_p.add_run("Tabela: Wykryte anomalie")
    _style_run(run, FONT_SIZE_CAPTION, bold=True)

    # Table
    table = doc.add_table(rows=1 + len(anomaly_rows), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    body.remove(table._tbl)
    idx = _body_index(body, caption_p._element)
    body.insert(idx + 1, table._tbl)

    columns = ["Kategoria", "Opis", "Dane"]
    for i, col_name in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        _style_cell_runs(cell, FONT_SIZE_TABLE, bold=True)

    for row_idx, row_data in enumerate(anomaly_rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < 3:
                cell = row.cells[col_idx]
                cell.text = _fix_orphans(str(cell_text))
                size = FONT_SIZE_TABLE_SMALL if col_idx == 2 else FONT_SIZE_TABLE
                _style_cell_runs(cell, size)

    _add_spacer(doc, body, table._tbl)
    log.debug("Inserted anomaly section with %d rows", len(anomaly_rows))


# ─── Table insertion ─────────────────────────────────────────────────────────

def _insert_tables(doc, body, table_data, selected_tables) -> None:
    """Insert data tables into rendered DOCX at section positions."""
    from docx.enum.table import WD_TABLE_ALIGNMENT

    for table_name in selected_tables:
        tdef = TABLE_DEFS.get(table_name)
        rows = table_data.get(table_name)
        if not tdef or not rows:
            continue

        insert_after = _find_section_last_element(body, tdef["section_heading"])
        if insert_after is None:
            log.warning("Section '%s' not found, skipping table '%s'",
                        tdef["section_heading"], table_name)
            continue

        caption_p = _insert_paragraph_after(doc, body, insert_after)
        run = caption_p.add_run(tdef["caption"])
        _style_run(run, FONT_SIZE_CAPTION, bold=True)

        columns = tdef["columns"]
        n_cols = len(columns)
        table = doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        body.remove(table._tbl)
        idx = _body_index(body, caption_p._element)
        body.insert(idx + 1, table._tbl)

        for i, col_name in enumerate(columns):
            cell = table.rows[0].cells[i]
            cell.text = col_name
            _style_cell_runs(cell, FONT_SIZE_TABLE, bold=True)

        for row_idx, row_data in enumerate(rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < n_cols:
                    cell = row.cells[col_idx]
                    cell.text = str(cell_text)
                    _style_cell_runs(cell, FONT_SIZE_TABLE)

        _add_spacer(doc, body, table._tbl)
        log.debug("Inserted table '%s' with %d rows", table_name, len(rows))


# ─── Chart image insertion ───────────────────────────────────────────────────

def _insert_chart_images(doc, body, chart_images: Dict[str, bytes], *, tpl_texts: Optional[Dict[str, str]] = None) -> None:
    """Insert chart PNG images into DOCX after relevant sections."""
    from docx.shared import Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    last_section5_elem = _find_section_last_element(body, "5.")
    last_section7_elem = _find_section_last_element(body, "7.")
    cursor_s5 = last_section5_elem
    cursor_s7 = last_section7_elem

    for chart_def in CHART_INSERTION_ORDER:
        key = chart_def["key"]
        img_bytes = chart_images.get(key)
        if not img_bytes:
            continue

        cursor = cursor_s7 if chart_def["section_heading"] == "7." else cursor_s5
        if cursor is None:
            log.warning("No insertion point for chart '%s'", key)
            continue

        # Section heading (numbered, e.g. "6. Rodzaj aktywności")
        if chart_def.get("new_section_heading"):
            heading_p = _insert_paragraph_after(doc, body, cursor)
            run = heading_p.add_run(chart_def["new_section_heading"])
            _style_run(run, 13, bold=True)
            # Apply heading style if possible
            try:
                heading_p.style = doc.styles['Heading 2']
            except Exception:
                pass  # fallback to bold formatting
            cursor = heading_p._element
        # Sub-heading (no number)
        elif chart_def.get("sub_heading"):
            heading_p = _insert_paragraph_after(doc, body, cursor)
            run = heading_p.add_run(chart_def["sub_heading"])
            _style_run(run, 11, bold=True)
            cursor = heading_p._element

        # Pre-text (may be overridden by user template)
        pre_text = chart_def.get("pre_text", "")
        if tpl_texts and key == "top_contacts" and "contacts_graph" in tpl_texts:
            pre_text = tpl_texts["contacts_graph"]
        elif tpl_texts and key == "activity" and "activity_distribution" in tpl_texts:
            pre_text = tpl_texts["activity_distribution"]
        if pre_text:
            text_p = _insert_paragraph_after(doc, body, cursor)
            run = text_p.add_run(_fix_orphans(pre_text))
            _style_run(run, FONT_SIZE_BODY)
            text_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cursor = text_p._element

        # Image
        try:
            img_stream = io.BytesIO(img_bytes)
            img_p = _insert_paragraph_after(doc, body, cursor)
            run = img_p.add_run()
            run.add_picture(img_stream, width=Mm(155))
            img_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cursor = img_p._element

            # Caption
            cap_p = _insert_paragraph_after(doc, body, cursor)
            run = cap_p.add_run(chart_def["caption"])
            _style_run(run, 8, italic=True, color=(100, 100, 100))
            cap_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cursor = cap_p._element

            log.debug("Inserted chart image '%s'", key)
        except Exception as e:
            log.warning("Failed to insert chart image '%s': %s", key, e)

        if chart_def["section_heading"] == "7.":
            cursor_s7 = cursor
        else:
            cursor_s5 = cursor


# ─── Bullet-list insertion ───────────────────────────────────────────────────

def _insert_bullet_list_after_text(doc, body, search_text: str, items: List[str]) -> None:
    """Insert bullet-point list (\u2022) after paragraph containing search_text."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    target = None
    for child in body:
        if child.tag == qn('w:p'):
            full_text = ''.join(child.itertext()).strip()
            if search_text in full_text:
                target = child
                break

    if target is None:
        log.debug("Text '%s' not found for bullet-list", search_text[:40])
        return

    cursor = target
    for item_text in items:
        # Ensure each bullet item ends with ";"
        text = _fix_orphans(item_text.rstrip(";").rstrip()) + ";"
        p = _insert_paragraph_after(doc, body, cursor)
        run = p.add_run(f"\u2022  {text}")
        _style_run(run, FONT_SIZE_BODY)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.8)
        cursor = p._element

    log.debug("Inserted %d bullet items after '%s'", len(items), search_text[:40])


# ─── Formatting fixes ───────────────────────────────────────────────────────

def _fix_formatting(doc) -> None:
    """Apply consistent formatting to all programmatically inserted paragraphs.

    - Body text: Calibri, justified, same size
    - Bullet lists: left-aligned
    - Fix orphan single characters (Polish typography)
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for para in doc.paragraphs:
        # Apply font to all runs that don't have a font set
        for run in para.runs:
            if run.font.name is None:
                run.font.name = FONT_NAME
            # Fix orphans in text
            if run.text:
                fixed = _fix_orphans(run.text)
                if fixed != run.text:
                    run.text = fixed


def _fix_orphans(text: str) -> str:
    """Replace spaces after Polish single-letter words with non-breaking spaces.

    Prevents 'w ', 'i ', 'z ', 'a ', 'o ', 'u ' from appearing at line ends.
    """
    if not text:
        return text
    return _ORPHAN_RE.sub('\\1\u00a0', text)


# ─── Page numbers ────────────────────────────────────────────────────────────

def _add_page_numbers(doc) -> None:
    """Add page numbers in footer: '<page>/<total>' right-aligned."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    section = doc.sections[-1]

    # Ensure footer exists
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing footer content
    for p in footer.paragraphs:
        for r in p.runs:
            r.text = ""

    # Create or use first paragraph
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add PAGE field
    run1 = para.add_run()
    run1.font.name = FONT_NAME
    run1.font.size = Pt(8)

    fld_page = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "/>'
    )
    run1._element.append(fld_page)

    # Add separator
    run2 = para.add_run("/")
    run2.font.name = FONT_NAME
    run2.font.size = Pt(8)

    # Add NUMPAGES field
    run3 = para.add_run()
    run3.font.name = FONT_NAME
    run3.font.size = Pt(8)

    fld_total = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "/>'
    )
    run3._element.append(fld_total)

    log.debug("Added page numbers to footer")


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _insert_paragraph_after(doc, body, after_elem):
    """Create a new paragraph and insert it after the given element."""
    p = doc.add_paragraph()
    body.remove(p._element)
    idx = _body_index(body, after_elem)
    body.insert(idx + 1, p._element)
    return p


def _style_run(run, size_pt: int, *, bold=False, italic=False, color=None):
    """Apply consistent styling to a run."""
    from docx.shared import Pt, RGBColor
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*(color or (0, 0, 0)))
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _style_cell_runs(cell, size_pt: int, *, bold=False):
    """Apply consistent styling to all runs in a table cell."""
    from docx.shared import Pt, RGBColor
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = FONT_NAME
            r.font.size = Pt(size_pt)
            r.font.color.rgb = RGBColor(0, 0, 0)
            if bold:
                r.bold = True


def _find_section_heading(body, heading_prefix: str):
    """Find the heading paragraph element for a section number."""
    from docx.oxml.ns import qn
    for child in body:
        if child.tag == qn('w:p'):
            pStyle = child.find(f'.//{qn("w:pStyle")}')
            is_heading = pStyle is not None and 'Heading' in pStyle.get(qn('w:val'), '')
            if is_heading:
                full_text = ''.join(child.itertext()).strip()
                if full_text.startswith(heading_prefix):
                    return child
    return None


def _find_section_last_element(body, heading_prefix: str):
    """Find the last body element in a section (before next heading)."""
    from docx.oxml.ns import qn
    found_heading = False
    last_elem = None
    for child in body:
        if child.tag == qn('w:p'):
            pStyle = child.find(f'.//{qn("w:pStyle")}')
            is_heading = pStyle is not None and 'Heading' in pStyle.get(qn('w:val'), '')
            if is_heading:
                full_text = ''.join(child.itertext()).strip()
                if full_text.startswith(heading_prefix):
                    found_heading = True
                    last_elem = child
                    continue
                elif found_heading:
                    break
        if found_heading:
            last_elem = child
    return last_elem


def _body_index(body, element) -> int:
    """Get index of element in body's children."""
    for i, child in enumerate(body):
        if child is element:
            return i
    return len(body) - 1


def _add_spacer(doc, body, after_elem) -> None:
    """Add a small spacer paragraph after an element."""
    from docx.shared import Pt
    spacer = _insert_paragraph_after(doc, body, after_elem)
    spacer.add_run("").font.size = Pt(6)


# ─── Table data builders ───────────────────────────────────────────────────

def build_table_data(gsm_data: dict, placeholders: Optional[Dict[str, Any]] = None) -> Dict[str, List[List[str]]]:
    """Build table row data from gsm_latest.json for all table types."""
    billing = gsm_data.get("billing", {})
    summary = billing.get("summary", {})
    analysis = billing.get("analysis", {})

    tables: Dict[str, List[List[str]]] = {}

    calls_out = summary.get("calls_out", 0) or 0
    calls_in = summary.get("calls_in", 0) or 0
    sms_out = summary.get("sms_out", 0) or 0
    sms_in = summary.get("sms_in", 0) or 0
    data_sessions = summary.get("data_sessions", 0) or 0
    total = summary.get("total_records", 0) or 0
    duration = summary.get("call_duration_seconds", 0) or 0
    unique = summary.get("unique_contacts", 0) or 0

    tables["stats"] = [
        ["Po\u0142\u0105czenia wychodz\u0105ce", str(calls_out)],
        ["Po\u0142\u0105czenia przychodz\u0105ce", str(calls_in)],
        ["SMS wychodz\u0105ce", str(sms_out)],
        ["SMS przychodz\u0105ce", str(sms_in)],
        ["Sesje transmisji danych", str(data_sessions)],
        ["\u0141\u0105czna liczba rekord\u00f3w", str(total)],
        ["Czas po\u0142\u0105cze\u0144 (s)", str(duration)],
        ["Unikalne kontakty", str(unique)],
    ]

    top_contacts = analysis.get("top_contacts", [])
    tables["contacts"] = []
    for c in top_contacts[:10]:
        tables["contacts"].append([
            c.get("number", "?"),
            str(c.get("total_interactions", 0)),
            str(c.get("calls_out", 0)),
            str(c.get("calls_in", 0)),
        ])

    if placeholders and "_anomaly_table_rows" in placeholders:
        tables["anomalies"] = placeholders["_anomaly_table_rows"]
    else:
        from backend.gsm.note_builder import _build_anomaly_table_rows
        anomalies = analysis.get("anomalies", [])
        tables["anomalies"] = _build_anomaly_table_rows(anomalies)

    locations = analysis.get("locations", [])
    total_loc_records = sum(loc.get("record_count", 0) for loc in locations) or 1
    tables["locations"] = []
    for loc in locations[:15]:
        location = loc.get("location", loc.get("address", "?"))
        records = str(loc.get("record_count", 0))
        pct = f"{loc.get('record_count', 0) / total_loc_records * 100:.1f}%"
        tables["locations"].append([location, records, pct])

    return tables


# ─── Template helpers ────────────────────────────────────────────────────────

def _flatten_for_template(placeholders: Dict[str, Any]) -> Dict[str, Any]:
    """Flatten nested dicts into dot-notation keys for docxtpl."""
    context: Dict[str, Any] = {}
    for key, value in placeholders.items():
        if key.startswith("_"):
            continue
        if isinstance(value, dict):
            context[key] = value
            for sub_key, sub_value in value.items():
                context[f"{key}_{sub_key}"] = sub_value
        else:
            context[key] = value
    return context


def _extract_template_texts(custom_template: Dict[str, Any]) -> Dict[str, str]:
    """Extract user-edited text blocks from a custom template.

    Maps text sections to named overrides based on their position relative
    to marker sections in the template.

    Returns a dict like:
      {
        "anomaly_intro": "...",
        "contacts_graph": "...",
        "activity_distribution": "...",
      }
    """
    texts: Dict[str, str] = {}
    sections = custom_template.get("sections", [])

    # Build mapping: text block ID → preceding/following marker keys
    for i, sec in enumerate(sections):
        if sec.get("type") != "text":
            continue

        content = sec.get("content", "").strip()
        if not content:
            continue

        # Strip HTML tags for the DOCX (we get contenteditable HTML)
        import re
        clean = re.sub(r'<[^>]+>', '', content).strip()
        if not clean:
            continue

        # Identify text block purpose by looking at surrounding markers
        next_marker = None
        for j in range(i + 1, len(sections)):
            if sections[j].get("type") == "marker":
                next_marker = sections[j].get("key", "")
                break

        prev_marker = None
        for j in range(i - 1, -1, -1):
            if sections[j].get("type") == "marker":
                prev_marker = sections[j].get("key", "")
                break

        # Map to known text override keys
        if next_marker == "table_anomalies" or prev_marker == "table_anomalies":
            texts["anomaly_intro"] = clean
        elif next_marker == "chart_top_contacts":
            texts["contacts_graph"] = clean
        elif next_marker == "chart_activity":
            texts["activity_distribution"] = clean

    return texts


def _apply_llm_overrides(placeholders: Dict[str, Any], overrides: Dict[str, str]) -> None:
    """Apply LLM-generated text overrides to placeholders dict."""
    for key, value in overrides.items():
        if key.startswith("_"):
            continue
        parts = key.split(".", 1)
        if len(parts) == 2:
            parent, child = parts
            if parent in placeholders and isinstance(placeholders[parent], dict):
                placeholders[parent][child] = value
            else:
                placeholders[parent] = {child: value}
        else:
            placeholders[key] = value


def get_default_template_path() -> Path:
    """Return path to the bundled GSM note template."""
    pkg_dir = Path(__file__).resolve().parent.parent.parent
    template = pkg_dir / "templates" / "gsm_note_template.docx"
    if template.exists():
        return template

    import os
    data_dir = Path(os.environ.get("AISTATEWEB_DATA_DIR", "data_www"))
    alt = data_dir / "templates" / "gsm_note_template.docx"
    if alt.exists():
        return alt

    raise FileNotFoundError(
        f"GSM note template not found. Expected at {template} or {alt}"
    )
