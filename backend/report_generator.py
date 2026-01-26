"""AISTATEweb — Analysis report generator.

This module is used by the Analysis tab to persist generated content as:
- Plain text (.txt)
- Markdown (.md)
- HTML (.html) – standalone file with embedded CSS
- DOC (.doc) – Word-compatible HTML (lightweight)
- DOCX (.docx) – best-effort Markdown -> Word conversion

It is intentionally conservative (no heavy rendering engines).
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple


class ReportSaveError(RuntimeError):
    pass


def _slugify(text: str, max_len: int = 80) -> str:
    """Filesystem-safe slug."""
    if not text:
        return "report"
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower().strip()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return (text or "report")[:max_len]


def _md_to_html(md_text: str) -> str:
    try:
        import markdown  # type: ignore

        return markdown.markdown(
            md_text,
            extensions=["tables", "fenced_code", "toc"],
            output_format="html5",
        )
    except Exception as e:  # pragma: no cover
        raise ReportSaveError(f"Markdown->HTML conversion failed: {e}")


def _html_wrap(body_html: str, title: str, meta: Dict[str, Any]) -> str:
    # Minimal, readable styling.
    generated_at = meta.get("generated_at") or datetime.now().strftime("%Y-%m-%d %H:%M")
    project_id = meta.get("project_id") or ""
    model = meta.get("model") or ""
    templates = meta.get("template_ids") or []
    if isinstance(templates, list):
        templates_s = ", ".join(str(x) for x in templates if str(x).strip())
    else:
        templates_s = str(templates)

    meta_bits = [
        f"<span><b>Generated:</b> {generated_at}</span>",
    ]
    if project_id:
        meta_bits.append(f"<span><b>Project:</b> {project_id}</span>")
    if model:
        meta_bits.append(f"<span><b>Model:</b> {model}</span>")
    if templates_s:
        meta_bits.append(f"<span><b>Prompts:</b> {templates_s}</span>")

    meta_html = "\n".join(meta_bits)

    return f"""<!doctype html>
<html lang=\"pl\">
<head>
  <meta charset=\"utf-8\"/>
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\"/>
  <title>{title}</title>
  <style>
    :root {{
      --fg: #111827;
      --muted: #6b7280;
      --bg: #ffffff;
      --card: #f9fafb;
      --border: #e5e7eb;
      --accent: #1e40af;
      --codebg: #0b1020;
      --codefg: #e5e7eb;
    }}
    html, body {{ background: var(--bg); color: var(--fg); }}
    body {{ font-family: -apple-system, BlinkMacSystemFont, \"Segoe UI\", Roboto, Arial, sans-serif; line-height: 1.6; }}
    .wrap {{ max-width: 980px; margin: 32px auto; padding: 0 18px 48px; }}
    h1, h2, h3, h4 {{ color: var(--accent); line-height: 1.25; }}
    h1 {{ border-bottom: 2px solid var(--border); padding-bottom: 10px; }}
    .meta {{ display: flex; flex-wrap: wrap; gap: 10px 18px; margin: 14px 0 22px; padding: 12px 14px; background: var(--card); border: 1px solid var(--border); border-radius: 10px; color: var(--muted); font-size: 14px; }}
    a {{ color: var(--accent); }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid var(--border); padding: 10px 12px; vertical-align: top; }}
    th {{ background: var(--card); text-align: left; }}
    pre {{ background: var(--codebg); color: var(--codefg); padding: 14px 16px; border-radius: 10px; overflow-x: auto; }}
    code {{ background: #f3f4f6; padding: 0 4px; border-radius: 6px; }}
    blockquote {{ border-left: 4px solid var(--border); margin: 12px 0; padding: 6px 12px; color: var(--muted); }}
    hr {{ border: none; border-top: 1px solid var(--border); margin: 22px 0; }}
  </style>
</head>
<body>
  <div class=\"wrap\">
    <h1>{title}</h1>
    <div class=\"meta\">{meta_html}</div>
    {body_html}
  </div>
</body>
</html>"""


def save_markdown(md_text: str, out_path: Path) -> None:
    out_path.write_text(md_text, encoding="utf-8")


def save_text(md_text: str, out_path: Path) -> None:
    """Save as plain text. (We keep original content; may contain markdown.)"""
    out_path.write_text(md_text, encoding="utf-8")


def save_html_from_markdown(md_text: str, out_path: Path, title: str, meta: Dict[str, Any]) -> None:
    body_html = _md_to_html(md_text)
    out_path.write_text(_html_wrap(body_html, title=title, meta=meta), encoding="utf-8")


def save_doc_from_markdown(md_text: str, out_path: Path, title: str, meta: Dict[str, Any]) -> None:
    """Save as .doc (Word-compatible HTML).

    DOC is implemented as HTML with a .doc extension. This is lightweight and
    opens correctly in Microsoft Word / LibreOffice.
    """
    body_html = _md_to_html(md_text)
    html = _html_wrap(body_html, title=title, meta=meta)

    # Add minimal Word namespaces for better compatibility (best-effort).
    if "xmlns:w=\"urn:schemas-microsoft-com:office:word\"" not in html:
        html = html.replace(
            "<html",
            "<html xmlns:o=\"urn:schemas-microsoft-com:office:office\" xmlns:w=\"urn:schemas-microsoft-com:office:word\"",
            1,
        )
    out_path.write_text(html, encoding="utf-8")


def _parse_md_table(block_lines: List[str]) -> Optional[List[List[str]]]:
    # Very small markdown table parser.
    rows: List[List[str]] = []
    for ln in block_lines:
        s = ln.strip()
        if not s or "|" not in s:
            return None
        # header separator like | --- | --- |
        if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", s):
            continue
        parts = [p.strip() for p in s.strip("|").split("|")]
        rows.append(parts)
    return rows if rows else None


def save_docx_from_markdown(md_text: str, out_path: Path, title: str, meta: Dict[str, Any]) -> None:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except Exception as e:  # pragma: no cover
        raise ReportSaveError(f"python-docx not available: {e}")

    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=1)

    # Meta block
    meta_lines: List[str] = []
    for k in ("generated_at", "project_id", "model"):
        if meta.get(k):
            meta_lines.append(f"{k}: {meta.get(k)}")
    if meta.get("template_ids"):
        meta_lines.append(f"template_ids: {meta.get('template_ids')}")
    if meta_lines:
        p = doc.add_paragraph(" | ".join(meta_lines))
        p.runs[0].italic = True

    # Best-effort markdown parsing
    lines = md_text.replace("\r\n", "\n").split("\n")
    in_code = False
    code_buf: List[str] = []
    table_buf: List[str] = []

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph("\n".join(code_buf))
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        code_buf = []

    def flush_table() -> None:
        nonlocal table_buf
        if not table_buf:
            return
        parsed = _parse_md_table(table_buf)
        table_buf = []
        if not parsed:
            return
        cols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=cols)
        for r_i, row in enumerate(parsed):
            for c_i in range(cols):
                tbl.cell(r_i, c_i).text = row[c_i] if c_i < len(row) else ""

    def add_paragraph(text: str) -> None:
        if not text.strip():
            return
        doc.add_paragraph(text.strip())

    # Merge wrapped lines into paragraphs (simple heuristic)
    para_buf: List[str] = []
    list_mode: Optional[str] = None  # "ul"|"ol"

    def flush_para() -> None:
        nonlocal para_buf
        if para_buf:
            add_paragraph(" ".join(s.strip() for s in para_buf if s.strip()))
            para_buf = []

    for ln in lines:
        s = ln.rstrip("\n")

        # Fenced code
        if s.strip().startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_para()
                in_code = True
            continue
        if in_code:
            code_buf.append(s)
            continue

        # Table blocks
        if s.strip().startswith("|") and "|" in s.strip()[1:]:
            flush_para()
            table_buf.append(s)
            continue
        else:
            flush_table()

        # Headings
        if s.startswith("#### "):
            flush_para()
            doc.add_heading(s[5:].strip(), level=4)
            continue
        if s.startswith("### "):
            flush_para()
            doc.add_heading(s[4:].strip(), level=3)
            continue
        if s.startswith("## "):
            flush_para()
            doc.add_heading(s[3:].strip(), level=2)
            continue
        if s.startswith("# "):
            flush_para()
            doc.add_heading(s[2:].strip(), level=1)
            continue

        # Lists
        m_ol = re.match(r"^(\d+)\.\s+(.*)$", s.strip())
        if s.strip().startswith(("- ", "* ", "• ")):
            flush_para()
            list_text = s.strip()[2:].strip() if s.strip().startswith(("- ", "* ")) else s.strip()[2:].strip()
            doc.add_paragraph(list_text, style="List Bullet")
            list_mode = "ul"
            continue
        if m_ol:
            flush_para()
            doc.add_paragraph(m_ol.group(2).strip(), style="List Number")
            list_mode = "ol"
            continue

        if not s.strip():
            flush_para()
            list_mode = None
            continue

        # Paragraph lines (join)
        para_buf.append(s)

    flush_para()
    flush_table()
    if in_code:
        flush_code()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


@dataclass
class SaveReportResult:
    filename: str
    format: str
    path: str
    download_url: str
    size_bytes: int


def save_report(
    *,
    reports_dir: Path,
    content: str,
    output_format: str,
    title: Optional[str] = None,
    template_ids: Optional[List[str]] = None,
    project_id: Optional[str] = None,
    model: Optional[str] = None,
) -> SaveReportResult:
    """Save analysis content to reports_dir in given format."""
    fmt = (output_format or "md").lower().strip()
    if fmt in ("markdown", "md"):
        fmt = "md"
    if fmt in ("text", "txt"):
        fmt = "txt"
    if fmt in ("word", "doc"):
        fmt = "doc"
    if fmt not in ("txt", "md", "html", "doc", "docx"):
        raise ReportSaveError("Unsupported output_format. Use: txt|md|html|doc|docx")

    reports_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    base = title or ("_".join(template_ids) if template_ids else "analysis")
    slug = _slugify(base)
    filename = f"{ts}_{slug}.{fmt}"
    out_path = reports_dir / filename

    meta = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "project_id": project_id or "",
        "model": model or "",
        "template_ids": template_ids or [],
    }
    safe_title = title or "Raport analizy"

    if fmt == "txt":
        save_text(content, out_path)
    elif fmt == "md":
        save_markdown(content, out_path)
    elif fmt == "html":
        save_html_from_markdown(content, out_path, title=safe_title, meta=meta)
    elif fmt == "doc":
        save_doc_from_markdown(content, out_path, title=safe_title, meta=meta)
    elif fmt == "docx":
        save_docx_from_markdown(content, out_path, title=safe_title, meta=meta)

    # Store metadata alongside (best-effort)
    try:
        (out_path.with_suffix(out_path.suffix + ".meta.json")).write_text(
            __import__("json").dumps(meta, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except Exception:
        pass

    size = out_path.stat().st_size if out_path.exists() else 0
    return SaveReportResult(
        filename=filename,
        format=fmt,
        path=str(out_path),
        download_url="",  # filled by API layer (depends on routes)
        size_bytes=size,
    )
