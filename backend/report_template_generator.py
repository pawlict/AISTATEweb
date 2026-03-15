"""
DOCX Template Generator — creates base templates with Content Controls (SDT).

Generates a .docx template file with:
- Header zone: Structured Document Tags for editable fields
  (INSTYTUCJA, ADRES, SYGNATURA, DATA, ANALITYK)
- Body zone: Jinja2 placeholder for report sections (docxtpl compatible)
- Footer zone: SDT for PODPIS, STOPKA

Content Controls (SDT) are OOXML elements compatible with both
Microsoft Word and ONLYOFFICE Docs.
"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# ─── SDT (Content Control) helpers ───────────────────────────────────────────

def _create_sdt_element(
    tag_name: str,
    placeholder_text: str,
    default_value: str = "",
    sdt_type: str = "text",
    font_name: str = "Calibri",
    font_size_pt: int = 11,
    bold: bool = False,
) -> etree._Element:
    """Create a Structured Document Tag (Content Control) XML element.

    Args:
        tag_name: SDT tag/alias identifier (e.g. "INSTYTUCJA")
        placeholder_text: Hint text shown when control is empty
        default_value: Initial text value
        sdt_type: "text" (plain), "richtext", or "date"
        font_name: Font family
        font_size_pt: Font size in points
        bold: Whether text is bold
    """
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    sdt = etree.SubElement(etree.Element("dummy"), qn("w:sdt"))

    # ── SDT Properties ──
    sdtPr = etree.SubElement(sdt, qn("w:sdtPr"))

    # Alias (display name)
    alias_el = etree.SubElement(sdtPr, qn("w:alias"))
    alias_el.set(qn("w:val"), tag_name)

    # Tag (programmatic identifier)
    tag_el = etree.SubElement(sdtPr, qn("w:tag"))
    tag_el.set(qn("w:val"), tag_name)

    # Lock: content control cannot be deleted, but content can be edited
    lock_el = etree.SubElement(sdtPr, qn("w:lock"))
    lock_el.set(qn("w:val"), "sdtLocked")

    # Placeholder
    ph_el = etree.SubElement(sdtPr, qn("w:placeholder"))
    docPart = etree.SubElement(ph_el, qn("w:docPart"))
    docPart.set(qn("w:val"), placeholder_text)

    # Show placeholder when empty
    show_empty = etree.SubElement(sdtPr, qn("w:showingPlcHdr"))

    # Type-specific properties
    if sdt_type == "date":
        date_el = etree.SubElement(sdtPr, qn("w:date"))
        if default_value:
            date_el.set(qn("w:fullDate"), default_value)
        fmt = etree.SubElement(date_el, qn("w:dateFormat"))
        fmt.set(qn("w:val"), "yyyy-MM-dd")
        lid = etree.SubElement(date_el, qn("w:lid"))
        lid.set(qn("w:val"), "pl-PL")

    # ── SDT Content ──
    sdtContent = etree.SubElement(sdt, qn("w:sdtContent"))
    p = etree.SubElement(sdtContent, qn("w:p"))

    # Run properties (font)
    r = etree.SubElement(p, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), str(font_size_pt * 2))  # Half-points
    szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(font_size_pt * 2))
    if bold:
        etree.SubElement(rPr, qn("w:b"))

    # Text
    t = etree.SubElement(r, qn("w:t"))
    t.set(qn("xml:space"), "preserve")
    t.text = default_value if default_value else placeholder_text

    # If showing placeholder, mark the run as placeholder style
    if not default_value:
        style_el = etree.SubElement(rPr, qn("w:rStyle"))
        style_el.set(qn("w:val"), "PlaceholderText")

    return sdt


def _find_and_update_sdt(doc_element, tag_name: str, new_value: str) -> bool:
    """Find an SDT by tag name in document XML and update its text value.

    Returns True if found and updated.
    """
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for sdt in doc_element.iter(qn("w:sdt")):
        sdtPr = sdt.find(qn("w:sdtPr"))
        if sdtPr is None:
            continue
        tag_el = sdtPr.find(qn("w:tag"))
        if tag_el is None:
            continue
        if tag_el.get(qn("w:val")) == tag_name:
            # Found it — update content
            sdtContent = sdt.find(qn("w:sdtContent"))
            if sdtContent is not None:
                # Find first run text
                for t_el in sdtContent.iter(qn("w:t")):
                    t_el.text = new_value
                    return True
                # If no text element, create one
                p = sdtContent.find(qn("w:p"))
                if p is None:
                    p = etree.SubElement(sdtContent, qn("w:p"))
                r = p.find(qn("w:r"))
                if r is None:
                    r = etree.SubElement(p, qn("w:r"))
                t = etree.SubElement(r, qn("w:t"))
                t.set(qn("xml:space"), "preserve")
                t.text = new_value
                return True

            # Remove placeholder indicator
            showPh = sdtPr.find(qn("w:showingPlcHdr"))
            if showPh is not None:
                sdtPr.remove(showPh)

            return True
    return False


# ─── Template generation ─────────────────────────────────────────────────────

def generate_gsm_template(
    placeholders: Optional[Dict[str, str]] = None,
) -> bytes:
    """Generate a GSM report DOCX template with Content Controls.

    Returns the template as bytes (ready to write to file).
    """
    return _generate_template("gsm", "Raport z analizy bilingu GSM", placeholders)


def generate_aml_template(
    placeholders: Optional[Dict[str, str]] = None,
) -> bytes:
    """Generate an AML report DOCX template with Content Controls."""
    return _generate_template("aml", "Raport z analizy AML", placeholders)


def _generate_template(
    report_type: str,
    title: str,
    placeholders: Optional[Dict[str, str]] = None,
) -> bytes:
    """Internal: generate base DOCX template."""
    if placeholders is None:
        placeholders = {}

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # ── Default style ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Body XML element ──
    body = doc.element.body

    # ── Header Content Controls ──
    # INSTYTUCJA (bold, 14pt)
    sdt_inst = _create_sdt_element(
        "INSTYTUCJA",
        "[Nazwa instytucji]",
        placeholders.get("INSTYTUCJA", ""),
        "richtext", bold=True, font_size_pt=14,
    )
    body.append(sdt_inst)

    # ADRES_INSTYTUCJI
    sdt_addr = _create_sdt_element(
        "ADRES_INSTYTUCJI",
        "[Adres instytucji]",
        placeholders.get("ADRES_INSTYTUCJI", ""),
        "richtext", font_size_pt=10,
    )
    body.append(sdt_addr)

    # Separator
    _add_paragraph(doc, "")

    # Report title
    p_title = doc.add_heading(title, level=1)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Separator line
    _add_paragraph(doc, "═" * 60)

    # SYGNATURA
    _add_label_with_sdt(body, "Sygnatura: ", "SYGNATURA",
                        "[Nr sprawy]", placeholders.get("SYGNATURA", ""))

    # DATA_RAPORTU
    today = datetime.now().strftime("%Y-%m-%d")
    _add_label_with_sdt(body, "Data sporządzenia: ", "DATA_RAPORTU",
                        "[Data]", placeholders.get("DATA_RAPORTU", today), "date")

    # ANALITYK
    _add_label_with_sdt(body, "Analityk: ", "ANALITYK",
                        "[Imię i nazwisko]", placeholders.get("ANALITYK", ""))

    # Separator
    _add_paragraph(doc, "═" * 60)
    _add_paragraph(doc, "")

    # ── Jinja2 section marker (for docxtpl) ──
    # This paragraph contains a Jinja2 tag that docxtpl will process
    p_marker = doc.add_paragraph()
    run = p_marker.add_run("{{REPORT_SECTIONS}}")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(9)
    run.font.italic = True

    _add_paragraph(doc, "")

    # ── Footer Content Controls ──
    _add_paragraph(doc, "═" * 60)

    # PODPIS
    _add_label_with_sdt(body, "Podpis: ", "PODPIS",
                        "[Podpis]", placeholders.get("PODPIS", ""), "richtext")

    _add_paragraph(doc, "")

    # STOPKA
    sdt_footer = _create_sdt_element(
        "STOPKA",
        "Wygenerowano w AISTATEweb",
        placeholders.get("STOPKA", "Wygenerowano w AISTATEweb"),
        "richtext", font_size_pt=9,
    )
    body.append(sdt_footer)

    # ── Save to bytes ──
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_paragraph(doc: Document, text: str, bold: bool = False,
                   size_pt: int = 11) -> None:
    """Add a simple text paragraph."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size_pt)
        run.bold = bold


def _add_label_with_sdt(
    body,
    label: str,
    tag_name: str,
    placeholder: str,
    default_value: str = "",
    sdt_type: str = "text",
) -> None:
    """Add a paragraph with a text label followed by an SDT content control."""
    # Create paragraph with label
    p = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r><w:rPr><w:b/></w:rPr><w:t xml:space="preserve">{label}</w:t></w:r>'
        f'</w:p>'
    )
    body.append(p)

    # Create inline SDT (appended as separate paragraph for simplicity)
    sdt = _create_sdt_element(tag_name, placeholder, default_value, sdt_type)
    body.append(sdt)


def update_sdt_values(doc_path: Path, values: Dict[str, str],
                      output_path: Optional[Path] = None) -> Path:
    """Open a DOCX, update Content Controls by tag, save.

    Args:
        doc_path: Path to existing DOCX with SDTs.
        values: Dict of tag_name → new_value.
        output_path: Where to save. Defaults to doc_path (overwrite).

    Returns:
        Path to saved file.
    """
    doc = Document(str(doc_path))
    body = doc.element.body

    for tag_name, new_value in values.items():
        _find_and_update_sdt(body, tag_name, new_value)

    save_to = output_path or doc_path
    doc.save(str(save_to))
    return save_to
