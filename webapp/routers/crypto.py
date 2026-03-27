"""Crypto transaction analysis API router.

Endpoints:
- POST /api/crypto/analyze        - upload CSV/JSON/XLSX file and run full pipeline
- GET  /api/crypto/detail         - get saved analysis detail for project
- GET  /api/crypto/report         - generate HTML report for project
- GET  /api/crypto/llm-stream     - SSE streaming LLM narrative analysis
"""

from __future__ import annotations

import json
import logging
import os
import tempfile
import traceback
from pathlib import Path
from typing import Any, Dict

from fastapi import APIRouter, File, Form, Query, Request, UploadFile
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import HTMLResponse, JSONResponse
from starlette.responses import StreamingResponse

log = logging.getLogger("aistate.api.crypto")

router = APIRouter()


def _data_dir() -> Path:
    return Path(os.environ.get("AISTATEWEB_DATA_DIR", "data_www"))


def _app_log(msg: str) -> None:
    try:
        from webapp.server import app_log
        app_log(msg)
    except Exception:
        pass


def _project_path(project_id: str) -> Path:
    return _data_dir() / "projects" / project_id


def _crypto_save_path(project_id: str) -> Path:
    p = _project_path(project_id) / "analysis"
    p.mkdir(parents=True, exist_ok=True)
    return p / "crypto_latest.json"


# ---------------------------------------------------------------------------
#  Upload & analyze
# ---------------------------------------------------------------------------

@router.post("/api/crypto/analyze")
async def crypto_analyze(
    request: Request,
    file: UploadFile = File(...),
    project_id: str = Form(""),
):
    """Upload a crypto CSV/JSON/XLSX file and run the full analysis pipeline."""
    try:
        filename = file.filename or "upload.csv"
        content = await file.read()
        file_size = len(content)
        _app_log(f"[Crypto] Upload: {filename} ({file_size} bytes)")
        log.info("Crypto analyze request: file=%s size=%d project=%s", filename, file_size, project_id or "(none)")

        # Save to temp file
        suffix = Path(filename).suffix or ".csv"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        # Run pipeline in threadpool (CPU-bound)
        from backend.crypto.pipeline import run_crypto_pipeline
        result = await run_in_threadpool(
            run_crypto_pipeline, tmp_path, project_id, filename
        )

        # Cleanup temp file
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

        if not result.get("ok"):
            errors = result.get("errors", ["Nieznany błąd parsowania"])
            _app_log(f"[Crypto] Parse FAILED: {filename} - {'; '.join(errors)}")
            log.warning("Crypto parse failed: file=%s errors=%s", filename, errors)
            return JSONResponse({
                "status": "error",
                "errors": errors,
            }, status_code=400)

        # Save result to project
        if project_id:
            try:
                save_path = _crypto_save_path(project_id)
                save_path.write_text(
                    json.dumps(result, ensure_ascii=False, indent=1),
                    encoding="utf-8",
                )
                _app_log(f"[Crypto] Saved analysis to {save_path.name}")
            except Exception as e:
                log.warning("Failed to save crypto analysis: %s", e)
                _app_log(f"[Crypto] Save ERROR: {e}")

        tx_count = result.get('tx_count', 0)
        risk = result.get('risk_score', 0)
        elapsed = result.get('elapsed_sec', 0)
        source = result.get('source', '?')
        chain = result.get('chain', '?')
        wallets = result.get('wallet_count', 0)
        alerts_count = len(result.get('alerts', []))

        _app_log(
            f"[Crypto] Done: {filename} - {source}/{chain}, "
            f"{tx_count} txs, {wallets} wallets, "
            f"risk={risk:.1f}/100, alerts={alerts_count}, "
            f"{elapsed:.1f}s"
        )
        log.info(
            "Crypto analyze done: file=%s source=%s chain=%s txs=%d wallets=%d "
            "risk=%.1f alerts=%d elapsed=%.2fs",
            filename, source, chain, tx_count, wallets, risk, alerts_count, elapsed,
        )

        return JSONResponse({"status": "ok", "result": result})

    except Exception as e:
        log.exception("Crypto analyze error")
        _app_log(f"[Crypto] ERROR: {type(e).__name__}: {e}")
        return JSONResponse({
            "status": "error",
            "errors": [str(e)],
        }, status_code=500)


# ---------------------------------------------------------------------------
#  Load saved analysis
# ---------------------------------------------------------------------------

@router.get("/api/crypto/detail")
async def crypto_detail(project_id: str = Query("")):
    """Load the last saved crypto analysis for a project."""
    if not project_id:
        return JSONResponse({"status": "error", "detail": "project_id required"}, status_code=400)

    save_path = _crypto_save_path(project_id)
    if not save_path.exists():
        log.debug("Crypto detail: no saved analysis for project=%s", project_id)
        return JSONResponse({"status": "ok", "result": None})

    try:
        data = json.loads(save_path.read_text(encoding="utf-8"))
        log.info("Crypto detail loaded: project=%s txs=%s", project_id, data.get("tx_count", "?"))
        return JSONResponse({"status": "ok", "result": data})
    except Exception as e:
        log.error("Crypto detail load error: project=%s error=%s", project_id, e)
        return JSONResponse({"status": "error", "detail": str(e)}, status_code=500)


# ---------------------------------------------------------------------------
#  LLM Narrative Analysis (SSE streaming)
# ---------------------------------------------------------------------------

@router.get("/api/crypto/config-stats")
async def crypto_config_stats():
    """Return counts of sanctioned addresses and known contracts."""
    import json as _json
    config_dir = Path(__file__).resolve().parent.parent.parent / "backend" / "crypto" / "config"
    ofac_count = 0
    contracts_count = 0
    try:
        sanc_path = config_dir / "sanctioned.json"
        if sanc_path.exists():
            data = _json.loads(sanc_path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                ofac_count = sum(len(v) if isinstance(v, list) else 1 for v in data.values())
            elif isinstance(data, list):
                ofac_count = len(data)
    except Exception as e:
        log.warning("Failed to load sanctioned.json: %s", e)
    try:
        cont_path = config_dir / "known_contracts.json"
        if cont_path.exists():
            data = _json.loads(cont_path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                contracts_count = sum(len(v) if isinstance(v, list) else 1 for v in data.values())
            elif isinstance(data, list):
                contracts_count = len(data)
    except Exception as e:
        log.warning("Failed to load known_contracts.json: %s", e)
    log.debug("Crypto config-stats: ofac=%d contracts=%d", ofac_count, contracts_count)
    return JSONResponse({
        "ofac_count": ofac_count,
        "contracts_count": contracts_count,
    })


# ---------------------------------------------------------------------------
#  OFAC Sanctioned Addresses management
# ---------------------------------------------------------------------------

@router.get("/api/crypto/ofac/stats")
async def crypto_ofac_stats():
    """Get detailed OFAC sanctioned addresses statistics."""
    try:
        from backend.crypto.ofac_importer import get_ofac_stats
        stats = await run_in_threadpool(get_ofac_stats)
        return JSONResponse({"status": "ok", **stats})
    except Exception as e:
        log.exception("OFAC stats error")
        return JSONResponse({"status": "error", "detail": str(e)}, status_code=500)


@router.post("/api/crypto/ofac/download")
async def crypto_ofac_download(request: Request):
    """Download OFAC SDN_ADVANCED.XML and import sanctioned crypto addresses."""
    try:
        body: Dict[str, Any] = {}
        try:
            body = await request.json()
        except Exception:
            pass

        url = body.get("url", "").strip()
        replace = body.get("replace", True)

        _app_log("[Crypto/OFAC] Starting online download...")
        log.info("OFAC download: url=%s replace=%s", url or "(default)", replace)

        from backend.crypto.ofac_importer import download_and_import, DEFAULT_SOURCE_URL
        result = await run_in_threadpool(
            download_and_import,
            url=url or DEFAULT_SOURCE_URL,
            replace=replace,
        )

        _app_log(
            f"[Crypto/OFAC] Download done: {result.get('total_addresses', 0)} addresses, "
            f"{result.get('entities', 0)} entities"
        )
        log.info("OFAC download done: %s", result)
        return JSONResponse(result)

    except Exception as e:
        log.exception("OFAC download error")
        _app_log(f"[Crypto/OFAC] Download ERROR: {type(e).__name__}: {e}")
        return JSONResponse({
            "status": "error",
            "detail": str(e),
        }, status_code=500)


@router.post("/api/crypto/ofac/import")
async def crypto_ofac_import(
    file: UploadFile = File(...),
):
    """Import OFAC sanctioned addresses from an uploaded XML file."""
    try:
        filename = file.filename or "sdn_advanced.xml"
        content = await file.read()
        file_size = len(content)
        _app_log(f"[Crypto/OFAC] Offline import: {filename} ({file_size} bytes)")
        log.info("OFAC import: file=%s size=%d", filename, file_size)

        # Save to temp
        import tempfile as _tempfile
        suffix = Path(filename).suffix or ".xml"
        with _tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)

        try:
            from backend.crypto.ofac_importer import import_from_file
            result = await run_in_threadpool(import_from_file, tmp_path, True)
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

        _app_log(
            f"[Crypto/OFAC] Import done: {result.get('total_addresses', 0)} addresses, "
            f"{result.get('entities', 0)} entities"
        )
        log.info("OFAC import done: %s", result)
        return JSONResponse(result)

    except Exception as e:
        log.exception("OFAC import error")
        _app_log(f"[Crypto/OFAC] Import ERROR: {type(e).__name__}: {e}")
        return JSONResponse({
            "status": "error",
            "detail": str(e),
        }, status_code=500)


@router.post("/api/crypto/ofac/clear")
async def crypto_ofac_clear():
    """Clear all OFAC-sourced addresses from sanctioned.json."""
    try:
        config_path = Path(__file__).resolve().parent.parent.parent / "backend" / "crypto" / "config" / "sanctioned.json"

        if config_path.exists():
            data = json.loads(config_path.read_text(encoding="utf-8"))
            # Remove OFAC-sourced addresses
            addresses = data.get("addresses", {})
            kept = {k: v for k, v in addresses.items()
                    if not (isinstance(v, dict) and "OFAC" in str(v.get("reason", "")))}
            data["addresses"] = kept
            data["_ofac_records"] = 0
            data["_last_update"] = None
            config_path.write_text(
                json.dumps(data, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            # Invalidate cache
            try:
                from backend.crypto.ofac_importer import _invalidate_cache
                _invalidate_cache()
            except Exception:
                pass

            removed = len(addresses) - len(kept)
            _app_log(f"[Crypto/OFAC] Cleared {removed} OFAC addresses, kept {len(kept)} manual")
            return JSONResponse({"status": "ok", "removed": removed, "kept": len(kept)})
        else:
            return JSONResponse({"status": "ok", "removed": 0, "kept": 0})
    except Exception as e:
        log.exception("OFAC clear error")
        return JSONResponse({"status": "error", "detail": str(e)}, status_code=500)


# ---------------------------------------------------------------------------
#  Transaction classifications (save/load per project)
# ---------------------------------------------------------------------------

def _classifications_path(project_id: str) -> Path:
    p = _project_path(project_id) / "analysis"
    p.mkdir(parents=True, exist_ok=True)
    return p / "crypto_classifications.json"


def _load_classifications(project_id: str) -> Dict[str, str]:
    path = _classifications_path(project_id)
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return {}
    return {}


def _save_classifications(project_id: str, cls_map: Dict[str, str]) -> None:
    path = _classifications_path(project_id)
    path.write_text(
        json.dumps(cls_map, ensure_ascii=False, indent=1),
        encoding="utf-8",
    )


@router.get("/api/crypto/classifications")
async def crypto_get_classifications(project_id: str = Query("")):
    """Load saved transaction classifications for a project."""
    if not project_id:
        return JSONResponse({"classifications": {}})
    try:
        cls_map = _load_classifications(project_id)
        return JSONResponse({"classifications": cls_map})
    except Exception as e:
        log.warning("Failed to load crypto classifications: %s", e)
        return JSONResponse({"classifications": {}})


@router.post("/api/crypto/classify")
async def crypto_classify(request: Request):
    """Save a single transaction classification."""
    try:
        body = await request.json()
        project_id = body.get("project_id", "")
        tx_id = body.get("tx_id", "")
        classification = body.get("classification", "")

        if not project_id or not tx_id or not classification:
            return JSONResponse(
                {"status": "error", "detail": "project_id, tx_id, classification required"},
                status_code=400,
            )

        cls_map = _load_classifications(project_id)
        cls_map[tx_id] = classification
        await run_in_threadpool(_save_classifications, project_id, cls_map)

        log.debug("Crypto classify: project=%s tx=%s cls=%s", project_id, tx_id[:16], classification)
        return JSONResponse({"status": "ok"})
    except Exception as e:
        log.exception("Crypto classify error")
        return JSONResponse({"status": "error", "detail": str(e)}, status_code=500)


@router.post("/api/crypto/classify-batch")
async def crypto_classify_batch(request: Request):
    """Save multiple transaction classifications at once."""
    try:
        body = await request.json()
        project_id = body.get("project_id", "")
        classifications = body.get("classifications", {})

        if not project_id or not classifications:
            return JSONResponse(
                {"status": "error", "detail": "project_id and classifications required"},
                status_code=400,
            )

        cls_map = _load_classifications(project_id)
        # Only add new classifications, don't override existing manual ones
        for tx_id, cls in classifications.items():
            if tx_id not in cls_map:
                cls_map[tx_id] = cls
        await run_in_threadpool(_save_classifications, project_id, cls_map)

        log.info("Crypto classify-batch: project=%s count=%d total=%d",
                 project_id, len(classifications), len(cls_map))
        return JSONResponse({"status": "ok", "saved": len(cls_map)})
    except Exception as e:
        log.exception("Crypto classify-batch error")
        return JSONResponse({"status": "error", "detail": str(e)}, status_code=500)


# ---------------------------------------------------------------------------
#  LLM Narrative Analysis (SSE streaming)
# ---------------------------------------------------------------------------

@router.get("/api/crypto/llm-stream")
async def crypto_llm_stream(
    model: str = Query(""),
    user_prompt: str = Query(""),
    project_id: str = Query(""),
):
    """SSE streaming LLM analysis for crypto transaction data."""

    # Load saved crypto analysis
    crypto_data: dict = {}
    if project_id:
        try:
            save_path = _crypto_save_path(project_id)
            if save_path.exists():
                crypto_data = json.loads(save_path.read_text(encoding="utf-8"))
        except Exception as e:
            log.warning("Crypto LLM: failed to load data: %s", e)

    llm_prompt = crypto_data.get("llm_prompt", "")
    if not llm_prompt:
        log.info("Crypto LLM: no data for project=%s", project_id)
        _app_log(f"[Crypto] LLM: brak danych do analizy (project={project_id})")
        async def err_gen():
            yield f"data: {json.dumps({'error': 'Brak danych kryptowalutowych do analizy. Najpierw zaimportuj plik CSV.', 'done': True})}\n\n"
        return StreamingResponse(err_gen(), media_type="text/event-stream")

    # Append user prompt if provided
    if user_prompt.strip():
        llm_prompt += f"\n\n## Dodatkowe polecenie użytkownika\n{user_prompt.strip()}"

    chosen_model = model.strip()
    _app_log(f"[Crypto] LLM stream start: model={chosen_model or '(default)'}, project={project_id}")
    log.info("Crypto LLM stream: model=%s project=%s prompt_len=%d", chosen_model or "(default)", project_id, len(llm_prompt))

    async def generate():
        try:
            from backend.ollama_client import OllamaClient, stream_analyze
            client = OllamaClient()
            system_msg = (
                "Jesteś ekspertem ds. analizy transakcji kryptowalutowych, "
                "blockchain forensics i przeciwdziałania praniu pieniędzy (AML). "
                "Odpowiadaj po polsku, profesjonalnym językiem. "
                "Formatuj odpowiedź używając nagłówków markdown (##, ###). "
                "Bądź precyzyjny, wskazuj konkretne adresy i kwoty."
            )
            chunk_count = 0
            async for chunk in stream_analyze(
                client, llm_prompt, model=chosen_model,
                system=system_msg,
                options={"temperature": 0.3, "num_ctx": 8192},
            ):
                chunk_count += 1
                yield f"data: {json.dumps({'chunk': chunk, 'done': False}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'chunk': '', 'done': True, 'chunks': chunk_count})}\n\n"
            _app_log(f"[Crypto] LLM stream done: {chunk_count} chunks")
            log.info("Crypto LLM stream done: chunks=%d", chunk_count)
        except Exception as e:
            log.exception("Crypto LLM stream error")
            _app_log(f"[Crypto] LLM ERROR: {type(e).__name__}: {e}")
            yield f"data: {json.dumps({'error': str(e), 'done': True})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


# ---------------------------------------------------------------------------
#  Report generation (HTML / TXT / DOCX)
# ---------------------------------------------------------------------------

@router.get("/api/crypto/report")
async def crypto_report(
    project_id: str = Query(""),
    formats: str = Query("html"),
):
    """Generate a report for the project's crypto analysis.

    Supported formats (comma-separated): html, txt, docx.
    Returns the first selected format inline; if DOCX, returns as download.
    """
    if not project_id:
        return JSONResponse({"status": "error", "detail": "project_id required"}, status_code=400)

    save_path = _crypto_save_path(project_id)
    if not save_path.exists():
        return JSONResponse({"status": "error", "detail": "No saved crypto analysis"}, status_code=404)

    try:
        result = json.loads(save_path.read_text(encoding="utf-8"))
    except Exception as e:
        return JSONResponse({"status": "error", "detail": str(e)}, status_code=500)

    fmt_list = [f.strip().lower() for f in formats.split(",") if f.strip()]
    if not fmt_list:
        fmt_list = ["html"]

    primary = fmt_list[0]

    if primary == "txt":
        txt = _build_crypto_report_txt(result)
        return StreamingResponse(
            iter([txt.encode("utf-8")]),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename=crypto_report.txt"},
        )

    if primary == "docx":
        try:
            docx_bytes = _build_crypto_report_docx(result)
            return StreamingResponse(
                iter([docx_bytes]),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=crypto_report.docx"},
            )
        except Exception:
            # Fallback: generate TXT if DOCX module unavailable
            txt = _build_crypto_report_txt(result)
            return StreamingResponse(
                iter([txt.encode("utf-8")]),
                media_type="text/plain; charset=utf-8",
                headers={"Content-Disposition": f"attachment; filename=crypto_report.txt"},
            )

    # Default: HTML - download as file
    html = _build_crypto_report_html(result)
    return StreamingResponse(
        iter([html.encode("utf-8")]),
        media_type="text/html; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=crypto_report.html"},
    )


def _resc(s) -> str:
    """HTML-escape a string."""
    return (str(s)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


_TX_TYPE_DESCRIPTIONS: Dict[str, str] = {
    "deposit": "Wpłata środków na giełdę z zewnętrznego portfela lub konta bankowego",
    "withdraw": "Wypłata środków z giełdy na zewnętrzny portfel lub konto bankowe",
    "withdrawal": "Wypłata środków z giełdy na zewnętrzny portfel lub konto bankowe",
    "buy": "Zakup kryptowaluty za inną walutę (fiat lub krypto)",
    "sell": "Sprzedaż kryptowaluty za inną walutę (fiat lub krypto)",
    "swap": "Wymiana jednej kryptowaluty na inną bezpośrednio",
    "trade": "Transakcja handlowa - kupno lub sprzedaż na rynku giełdowym",
    "transfer": "Przeniesienie środków między własnymi kontami/portfelami na giełdzie",
    "send": "Wysłanie środków do innego użytkownika lub na zewnętrzny adres",
    "receive": "Otrzymanie środków od innego użytkownika lub z zewnętrznego adresu",
    "staking": "Zablokowanie kryptowaluty w celu uzyskania nagród (oprocentowanie)",
    "staking_reward": "Nagroda otrzymana za udział w stakingu (oprocentowanie krypto)",
    "learn_reward": "Nagroda za ukończenie kursu edukacyjnego na platformie",
    "unstaking": "Odblokowanie wcześniej zablokowanych środków ze stakingu",
    "earn": "Program oszczędnościowy/inwestycyjny - odsetki od zdeponowanych środków",
    "distribution": "Dystrybucja tokenów - airdrop, nagroda lub podział zysku",
    "airdrop": "Darmowe tokeny otrzymane w ramach promocji lub dystrybucji projektu",
    "fee": "Opłata transakcyjna pobrana przez giełdę lub sieć blockchain",
    "commission": "Prowizja pobrana przez giełdę za wykonanie transakcji",
    "funding": "Opłata za utrzymanie pozycji futures/margin (funding rate)",
    "futures": "Transakcja na kontrakcie terminowym (futures) - instrumenty pochodne",
    "margin": "Transakcja z dźwignią finansową (pożyczone środki)",
    "liquidation": "Przymusowe zamknięcie pozycji z powodu niewystarczającego zabezpieczenia",
    "convert": "Konwersja jednej kryptowaluty na inną po aktualnym kursie",
    "p2p": "Transakcja peer-to-peer - bezpośrednia wymiana między użytkownikami",
    "otc": "Transakcja OTC (Over-The-Counter) - poza rynkiem giełdowym, zwykle duże kwoty",
    "nft": "Transakcja związana z NFT (Non-Fungible Token) - unikalne tokeny cyfrowe",
    "mint": "Utworzenie nowego tokena lub NFT na blockchainie",
    "burn": "Trwałe zniszczenie/usunięcie tokenów z obiegu (zmniejszenie podaży)",
    "bridge": "Transfer kryptowaluty między różnymi blockchainami przez most (bridge)",
    "wrap": "Zamiana tokena na jego opakowaną wersję kompatybilną z innym blockchainem",
    "unwrap": "Zamiana opakowanego tokena z powrotem na oryginał",
    "loan": "Pożyczka kryptowalutowa - wypożyczenie lub zaciągnięcie pożyczki",
    "repayment": "Spłata pożyczki kryptowalutowej",
    "collateral": "Środki zablokowane jako zabezpieczenie pożyczki lub pozycji margin",
    "savings": "Środki ulokowane w programie oszczędnościowym giełdy",
    "launchpad": "Udział w sprzedaży nowego tokena (IEO/IDO) na platformie giełdy",
    "referral": "Premia/nagroda za polecenie giełdy innym użytkownikom",
    "cashback": "Zwrot części opłaty transakcyjnej lub zakupu",
    "dust_conversion": "Zamiana niewielkich resztek tokenów (dust) na jedną kryptowalutę",
    "incoming": "Transakcja przychodząca - środki otrzymane na portfel",
    "outgoing": "Transakcja wychodząca - środki wysłane z portfela",
    "contract_call": "Wywołanie smart kontraktu na blockchainie (np. interakcja z DeFi)",
    "approval": "Udzielenie zgody smart kontraktowi na zarządzanie tokenami",
    "self_transfer": "Transfer do samego siebie - przeniesienie między własnymi adresami",
}


def _build_crypto_report_html(r: Dict[str, Any]) -> str:
    """Build a standalone HTML report with logical section ordering.

    Sections:
      I.   Identyfikacja - account info, user IDs
      II.  Podsumowanie ogólne - stats, tokens, date range
      III. Profil zachowania - user behavior profiling
      IV.  Ocena ryzyka AML - risk score, reasons, alerts
      V.   Portfel tokenów - token breakdown
      VI.  Kontrahenci i transfery - counterparties, pay C2C, phones
      VII. Adresy on-chain - external src/dst, wallets, deposit addresses
      VIII.Analiza kryminalistyczna - pass-through, privacy coins, mining
      IX.  Bezpieczeństwo konta - access logs, devices, card timeline
      X.   Transakcje - full transaction list
    """
    from datetime import datetime

    filename = r.get("filename", "")
    date_from = (r.get("date_from", "") or "")[:10]
    date_to = (r.get("date_to", "") or "")[:10]
    risk_score = r.get("risk_score", 0)
    tx_count = r.get("tx_count", 0)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    em = r.get("exchange_meta", {})
    fr = r.get("forensic_report", {}) or {}
    ai = fr.get("account_info", {}) or {}
    bs = r.get("binance_summary", {}) or {}
    source = r.get("source", "?")

    if risk_score >= 70:
        rl, rc = "KRYTYCZNE", "#dc2626"
    elif risk_score >= 50:
        rl, rc = "WYSOKIE", "#f97316"
    elif risk_score >= 25:
        rl, rc = "SREDNIE", "#eab308"
    else:
        rl, rc = "NISKIE", "#22c55e"

    sn = 0  # section number
    sections = []

    # Section descriptions for analysts
    _SEC_DESC = {
        "podsumowanie": (
            "Podstawowe statystyki analizowanego konta kryptowalutowego. Sekcja zawiera informacje o łącznej liczbie transakcji, "
            "zakresie dat aktywności, listę obsługiwanych tokenów (kryptowalut) oraz źródło danych (np. nazwa giełdy, eksport z blockchainu). "
            "Te dane pozwalają szybko ocenić skalę aktywności - czy mamy do czynienia z kontem o kilku transakcjach, "
            "czy z intensywnie używanym rachunkiem z tysiącami operacji rozłożonych na lata."
        ),
        "profil": (
            "Na podstawie wzorców transakcji system automatycznie rozpoznaje typ użytkownika - np. inwestor długoterminowy (kupuje i trzyma), "
            "aktywny trader (częste kupna/sprzedaże), arbitrażysta (wykorzystuje różnice cenowe między giełdami) czy podmiot tranzytowy "
            "(środki szybko przechodzą przez konto). Profil jest wyznaczany algorytmicznie na podstawie częstotliwości transakcji, "
            "czasu trzymania tokenów, rodzajów operacji i ich wolumenów. Wynik zawiera procentowy poziom pewności oraz listę powodów, "
            "dla których dany profil został przypisany. Jeden użytkownik może wykazywać cechy kilku profili jednocześnie."
        ),
        "ryzyko": (
            "Kluczowa sekcja z punktu widzenia przeciwdziałania praniu pieniędzy (AML - Anti-Money Laundering). "
            "System analizuje wszystkie transakcje i przypisuje zagregowany wynik ryzyka w skali 0–100 punktów, "
            "gdzie 0 oznacza brak podejrzanych zachowań, a 100 - bardzo wysokie ryzyko. Ocena jest budowana z wielu czynników, m.in.: "
            "structuring (dzielenie kwot na mniejsze, aby ominąć progi raportowe), privacy coins (użycie kryptowalut utrudniających śledzenie, "
            "np. Monero, Zcash), szybkie przeloty (środki wpływające i wypływające w ciągu godzin), kontrahenci wysokiego ryzyka "
            "(transakcje z adresami powiązanymi z mikserami, sankcjami lub nielegalną działalnością). "
            "Poniżej wyniku punktowego znajduje się lista konkretnych czynników, które wpłynęły na ocenę, oraz alerty wygenerowane przez silnik reguł."
        ),
        "portfel": (
            "Zestawienie wszystkich kryptowalut (tokenów), które pojawiły się na analizowanym koncie. "
            "Wpływy i wypływy to łączna wartość otrzymana i wysłana w danym tokenie; saldo netto to różnica (dodatnie = więcej otrzymano niż wysłano). "
            "Kategoria klasyfikuje token (np. L1/infrastructure, stablecoin, payments/transfers, memecoin, privacy coin) - "
            "pomaga ocenić charakter aktywności. Alert wskazuje poziom ryzyka tokena (NORMAL, MEDIUM, HIGH, CRITICAL); "
            "token oznaczony jako HIGH/CRITICAL wymaga szczególnej uwagi analityka."
        ),
        "wykresy": (
            "Graficzna prezentacja danych ułatwiająca szybkie wychwycenie trendów i anomalii. "
            "Wykres salda w czasie pokazuje jak zmieniało się saldo każdego tokena na przestrzeni analizowanego okresu - "
            "gwałtowne skoki lub spadki mogą wskazywać na duże jednorazowe wpłaty/wypłaty. "
            "Gdy tokeny mają drastycznie różne wartości, wykres stosuje normalizację logarytmiczną (wszystkie linie skalowane procentowo). "
            "Graf przepływu transakcji przedstawia powiązania między adresami i kontrahentami - węzły to adresy/portfele, "
            "strzałki to kierunki przepływu środków, kolor odzwierciedla poziom ryzyka."
        ),
        "kontrahenci": (
            "Lista wszystkich podmiotów, z którymi analizowane konto wymieniało środki. Kontrahenci są identyfikowani na podstawie "
            "wewnętrznych ID użytkowników giełdy, adresów portfeli Binance Pay (C2C) oraz numerów telefonów wykrytych w danych transakcyjnych. "
            "Dla każdego kontrahenta podana jest liczba transakcji, łączne wpływy i wypływy, używane tokeny oraz okres współpracy. "
            "Pozwala zidentyfikować najaktywniejszych partnerów handlowych, wykryć nietypowe powiązania "
            "(np. jeden kontrahent odpowiada za większość obrotu) i powiązać numery telefonów z konkretnymi transakcjami."
        ),
        "adresy": (
            "Zestawienie adresów blockchain (ciągów znaków identyfikujących portfele w sieci) powiązanych z analizowanym kontem. "
            "Adresy depozytowe to portfele przypisane do użytkownika na giełdzie, na które wpływają środki z zewnątrz. "
            "Adresy zewnętrzne to portfele spoza giełdy, z których użytkownik otrzymywał środki lub na które je wysyłał. "
            "Adresy dwukierunkowe (zarówno wpłaty jak wypłaty) mogą wskazywać na własne portfele użytkownika poza giełdą "
            "lub na bliskie relacje handlowe."
        ),
        "kryminalistyczna": (
            "Zaawansowane algorytmy wykrywania podejrzanych wzorców. "
            "Przeloty tranzytowe (pass-through) to pary transakcji, gdzie depozyt i wypłata podobnej kwoty następują w ciągu 24 godzin - "
            "klasyczny wzorzec konta tranzytowego, przez które środki 'przelatują' bez dłuższego zatrzymania. "
            "Kryptowaluty prywatności (Monero, Zcash) utrudniają lub uniemożliwiają śledzenie nadawcy, odbiorcy i kwoty transakcji. "
            "Wzorce górnicze to regularne, małe wpływy z tego samego adresu mogące wskazywać na dochody z kopania kryptowalut. "
            "Handel z dźwignią (margin) obejmuje zlecenia margin/futures - intensywny handel lewarowany może wskazywać na spekulację wysokiego ryzyka."
        ),
        "bezpieczenstwo": (
            "Dane techniczne dotyczące sposobu korzystania z konta. "
            "Logi dostępu to historia logowań z adresami IP i statusami - wiele nieudanych prób może wskazywać na atak brute-force lub przejęcie konta. "
            "Odciski urządzeń zawierają informacje o przeglądarkach i systemach operacyjnych - nagła zmiana urządzenia to sygnał ostrzegawczy. "
            "Geolokalizacja pokazuje kraje logowań - logowania z wielu krajów w krótkim czasie mogą wskazywać na użycie VPN/proxy lub współdzielenie konta. "
            "Timeline kart płatniczych prezentuje aktywność kart powiązanych z kontem."
        ),
        "czasowa": (
            "Rozkład transakcji w podziale na godziny dnia (0–23) i dni tygodnia. "
            "Aktywność automatyczna - boty handlowe działają 24/7 z równomiernym rozkładem; człowiek ma naturalną przerwę nocną. "
            "Koordynacja - transakcje konsekwentnie realizowane w wąskim oknie czasowym mogą wskazywać na skoordynowane działania. "
            "Nietypowe wzorce - np. wyłącznie nocna aktywność może sugerować, że rzeczywisty użytkownik jest w innej strefie czasowej."
        ),
        "konwersje": (
            "Sekwencje szybkich zamian jednego tokena na drugi, np. PLN → BTC → XMR → USDT → EUR. "
            "Wieloetapowe konwersje realizowane w krótkim czasie to klasyczna technika layeringu - "
            "tworzenia kolejnych warstw transakcji, aby utrudnić prześledzenie pochodzenia środków."
        ),
        "structuring": (
            "Structuring (smurfing) polega na dzieleniu dużej kwoty na wiele mniejszych transakcji - "
            "często tuż poniżej progów raportowych (np. 15 000 EUR w UE, 10 000 USD w USA). "
            "Celem jest uniknięcie automatycznego zgłoszenia transakcji do jednostki analityki finansowej. "
            "System wykrywa grupy transakcji o zbliżonych kwotach, realizowanych w krótkim czasie, na ten sam lub powiązane konta."
        ),
        "wash": (
            "Wash trading to transakcje, w których ten sam podmiot (lub podmioty powiązane) występuje jednocześnie jako kupujący i sprzedający. "
            "Celem jest sztuczne generowanie wolumenu obrotu - np. aby token wyglądał na bardziej płynny niż jest w rzeczywistości, "
            "lub aby wygenerować fałszywą historię handlową. "
            "System wykrywa pary transakcji kupno↔sprzedaż tego samego tokena w krótkim czasie z podobnymi kwotami."
        ),
        "fiat": (
            "Punkty styku między tradycyjnym systemem finansowym a światem kryptowalut. "
            "On-ramp (wpłata fiat) to wpłaty z konta bankowego na giełdę; off-ramp (wypłata fiat) to wypłaty z giełdy na konto bankowe. "
            "Duże wpłaty fiat wymagają weryfikacji źródła środków; duże wypłaty mogą wskazywać na realizację zysków lub wyprowadzanie środków. "
            "Asymetria - np. duże wpłaty fiat bez odpowiadających wypłat krypto (lub odwrotnie) - może wskazywać na pranie pieniędzy."
        ),
        "p2p": (
            "Transakcje peer-to-peer to bezpośrednia wymiana kryptowalut między użytkownikami, często z pominięciem standardowej książki zleceń giełdy. "
            "Handel P2P jest popularny w krajach z ograniczeniami bankowymi i bywa wykorzystywany do omijania procedur KYC "
            "(Know Your Customer - weryfikacja tożsamości) oraz AML. "
            "Powtarzający się kontrahenci z dużymi wolumenami mogą wskazywać na niezarejestrowaną działalność wymiany walut."
        ),
        "velocity": (
            "Analiza tego, jak szybko środki przechodzą przez konto. "
            "Średni czas trzymania to ile czasu token pozostaje na koncie między wpłatą a wypłatą - "
            "normalni użytkownicy trzymają kryptowaluty dni lub tygodnie; konta tranzytowe - minuty lub godziny. "
            "Wskaźniki hot wallet oznaczają tokeny ze średnim czasem trzymania poniżej 1 godziny - "
            "zachowanie typowe dla portfeli gorących giełd lub kont używanych do przepuszczania środków. "
            "Zbliżone wartości wpłat i wypłat przy krótkim trzymaniu to silny sygnał konta tranzytowego."
        ),
        "fees": (
            "Zestawienie opłat transakcyjnych (gas fees, withdrawal fees, trading fees) w podziale na tokeny. "
            "Pozwala oszacować koszty aktywności na koncie, wykryć transakcje z nietypowo wysokimi opłatami "
            "(celowe podnoszenie opłat może wskazywać na front-running - wyprzedzanie cudzych transakcji w bloku), "
            "oraz zidentyfikować dominującą sieć blockchain na podstawie tokena opłat (np. ETH = Ethereum, BNB = BSC)."
        ),
        "sieci": (
            "Statystyki wykorzystania różnych sieci blockchain (np. Ethereum, Binance Smart Chain, Tron, Polygon). "
            "Dla każdej sieci podana jest liczba transakcji i wolumen. Pozwala określić, które sieci użytkownik preferuje "
            "(np. Tron jest popularny do tanich transferów USDT), wykryć bridge'owanie - przenoszenie środków między łańcuchami, "
            "co może utrudniać śledzenie - oraz ocenić, czy użytkownik korzysta z sieci o niższym poziomie monitoringu AML."
        ),
        "ext_security": (
            "Pogłębiona analiza techniczna sesji i urządzeń. "
            "Anomalie logowań to wykryte nietypowe wzorce dostępu (logowanie z nowego kraju, zmiana fingerprintu urządzenia). "
            "VPN/Proxy - wykryte logowania z adresów IP znanych dostawców VPN lub sieci Tor; "
            "ukrywanie lokalizacji samo w sobie nie jest nielegalne, ale w kontekście AML jest czynnikiem ryzyka. "
            "Geolokalizacja kart - porównanie lokalizacji transakcji kartowych z lokalizacją logowań; "
            "rozbieżności mogą wskazywać na kradzież tożsamości lub cloning karty."
        ),
        "transakcje": (
            "Kompletna lista wszystkich transakcji - baza dowodowa dla całej analizy. "
            "Każdy wiersz zawiera datę, typ operacji (deposit, withdrawal, trade, swap, transfer itp.), "
            "token i kwotę, adresy nadawcy i odbiorcy, automatycznie przypisany poziom ryzyka dla tej transakcji "
            "oraz tagi klasyfikacyjne (np. suspicious, high_value, privacy_coin). "
            "Tabela pozwala analitykowi na ręczną weryfikację dowolnej transakcji wychwycone przez algorytmy w sekcjach powyżej."
        ),
    }

    def _sdesc(key: str) -> str:
        """Return HTML paragraph with section description, or empty string."""
        txt = _SEC_DESC.get(key, "")
        return f'<p class="section-desc">{txt}</p>' if txt else ""

    # ── I. Identyfikacja ──
    sn += 1
    meta = r.get("metadata", {}) or {}
    id_rows = []
    for label, val in [
        ("Właściciel konta", ai.get("holder_name") or meta.get("account_holder")),
        ("Imię", ai.get("first_name")),
        ("Nazwisko", ai.get("last_name")),
        ("User ID", ai.get("user_id")),
        ("Email", ai.get("email")),
        ("Telefon", ai.get("phone")),
        ("Data urodzenia", ai.get("date_of_birth")),
        ("Płeć", ai.get("gender")),
        ("Kraj zamieszkania", ai.get("country") or meta.get("country")),
        ("Narodowość", ai.get("nationality")),
        ("Adres zamieszkania", ai.get("physical_address") or meta.get("street")),
        ("Miasto", ai.get("city") or meta.get("city")),
        ("Województwo/Stan", ai.get("state")),
        ("Kod pocztowy", ai.get("zip_code") or meta.get("postal_code")),
        ("Poziom KYC", ai.get("kyc_level")),
        ("Poziom VIP", ai.get("vip_level")),
        ("Data rejestracji", ai.get("registration_date")),
        ("Status konta", ai.get("account_status")),
        ("Typ dokumentu", ai.get("id_type")),
        ("Nr dokumentu", ai.get("id_number")),
        ("ID polecającego", ai.get("referral_id")),
        ("ID agenta", ai.get("agent_id")),
        ("Sub-konto", ai.get("sub_account")),
        ("Margin", ai.get("margin_enabled")),
        ("Futures", ai.get("futures_enabled")),
        ("API Trading", ai.get("api_trading")),
        ("Kod anti-phishing", ai.get("anti_phishing_code")),
        ("Platforma", em.get("exchange_name") or source),
        ("Plik źródłowy", filename),
    ]:
        if val:
            id_rows.append(f"<tr><th>{_resc(label)}</th><td>{_resc(val)}</td></tr>")

    # User IDs across sheets
    uids = fr.get("user_ids_in_file", {})
    uid_html = ""
    if uids:
        uid_rows = "".join(f"<tr><td>{_resc(sh)}</td><td><code>{_resc(', '.join(ids))}</code></td></tr>"
                           for sh, ids in uids.items())
        uid_html = f'<h3>User IDs w poszczególnych arkuszach</h3><table class="data-table"><thead><tr><th>Arkusz</th><th>User ID</th></tr></thead><tbody>{uid_rows}</tbody></table>'

    sections.append(f'<h2>{sn}. Identyfikacja podmiotu</h2><table class="info-table">{"".join(id_rows)}</table>{uid_html}')

    # ── II. Podsumowanie ogólne ──
    sn += 1
    sum_rows = [
        f"<tr><th>Okres analizy</th><td>{_resc(date_from)} - {_resc(date_to)}</td></tr>",
        f"<tr><th>Łączna liczba transakcji</th><td>{tx_count}</td></tr>",
        f"<tr><th>Portfele / adresy</th><td>{r.get('wallet_count', 0)}</td></tr>",
        f"<tr><th>Kontrahenci</th><td>{r.get('counterparty_count', 0)}</td></tr>",
        f"<tr><th>Unikalne tokeny</th><td>{len(r.get('tokens', {}))}</td></tr>",
        f"<tr><th>Ryzyko AML</th><td style='color:{rc};font-weight:bold'>{risk_score:.1f}/100 ({rl})</td></tr>",
    ]
    if em.get("crypto_tokens"):
        sum_rows.append(f"<tr><th>Tokeny krypto</th><td>{_resc(', '.join(em['crypto_tokens']))}</td></tr>")
    if em.get("fiat_tokens"):
        sum_rows.append(f"<tr><th>Waluty fiat</th><td>{_resc(', '.join(em['fiat_tokens']))}</td></tr>")
    if em.get("account_types"):
        sum_rows.append(f"<tr><th>Typy kont</th><td>{_resc(', '.join(em['account_types']))}</td></tr>")

    # Fiat in/out
    fiat_in = bs.get("fiat_in", {})
    fiat_out = bs.get("fiat_out", {})
    for cur, amt in fiat_in.items():
        sum_rows.append(f"<tr><th>Wpłaty fiat ({_resc(cur)})</th><td>{amt:,.2f}</td></tr>")
    for cur, amt in fiat_out.items():
        sum_rows.append(f"<tr><th>Wypłaty fiat ({_resc(cur)})</th><td>{amt:,.2f}</td></tr>")

    sections.append(f'<h2>{sn}. Podsumowanie ogólne</h2>{_sdesc("podsumowanie")}<table class="info-table">{"".join(sum_rows)}</table>')

    # ── III. Profil zachowania ──
    bp = r.get("behavior_profile", {})
    if bp and bp.get("profiles"):
        sn += 1
        ph = ""
        for p in bp["profiles"][:5]:
            if p["score"] < 15:
                continue
            reasons = ""
            if p.get("reasons"):
                reasons = "<ul>" + "".join(f"<li>{_resc(rr)}</li>" for rr in p["reasons"]) + "</ul>"
            bc = '#22c55e' if p['score'] >= 50 else '#eab308' if p['score'] >= 30 else '#94a3b8'
            ph += f'<div class="profile-card" style="border-left:4px solid {bc}"><strong>{_resc(p["icon"])} {_resc(p["label"])}</strong> - {p["score"]}%<div class="desc">{_resc(p["desc"])}</div>{reasons}</div>'
        sections.append(f"<h2>{sn}. Profil zachowania użytkownika</h2>{_sdesc('profil')}{ph}")

    # ── IV. Ocena ryzyka AML ──
    sn += 1
    risk_html = f'<div style="font-size:18px;font-weight:bold;color:{rc};margin-bottom:12px">{risk_score:.1f}/100 - {rl}</div>'
    risk_reasons = r.get("risk_reasons", [])
    if risk_reasons:
        risk_html += "<h3>Czynniki ryzyka</h3><ul>" + "".join(f"<li>{_resc(rr)}</li>" for rr in risk_reasons) + "</ul>"
    alerts = r.get("alerts", [])
    if alerts:
        alert_items = "".join(f"<li><strong>{_resc(a.get('type', ''))}</strong>: {_resc(a.get('message', ''))}</li>" for a in alerts)
        risk_html += f"<h3>Alerty ({len(alerts)})</h3><ul>{alert_items}</ul>"
    sections.append(f"<h2>{sn}. Ocena ryzyka AML</h2>{_sdesc('ryzyko')}{risk_html}")

    # ── V. Portfel tokenów ──
    tokens = r.get("tokens", {})
    tc = r.get("token_classification", {}) or {}
    if tokens:
        sn += 1
        alert_colors = {"CRITICAL": "#dc2626", "HIGH": "#f97316", "MEDIUM": "#eab308", "NORMAL": "#22c55e"}
        rows = ""
        for tok, s in sorted(tokens.items(), key=lambda x: x[1].get("count", 0), reverse=True):
            net = (s.get("received", 0) or 0) - (s.get("sent", 0) or 0)
            nc = "#22c55e" if net >= 0 else "#dc2626"
            info = tc.get(tok, {})
            name = info.get("name", "")
            cat = info.get("category", "")
            alert = info.get("alert_level", "NORMAL")
            ac = alert_colors.get(alert, "#94a3b8")
            desc = _resc(info.get("description", ""))
            risk_note = _resc(info.get("risk_note", ""))
            tooltip = f' title="{risk_note}"' if risk_note else ""
            rows += (f"<tr><td style='font-weight:600'>{_resc(tok)}</td>"
                     f"<td>{_resc(name)}</td>"
                     f"<td><span style='font-size:11px;padding:1px 6px;border-radius:3px;background:#f1f5f9'>{_resc(cat)}</span></td>"
                     f"<td class='num'>{s.get('received', 0):.4f}</td>"
                     f"<td class='num'>{s.get('sent', 0):.4f}</td>"
                     f"<td class='num' style='color:{nc}'>{net:.4f}</td>"
                     f"<td>{s.get('count', 0)}</td>"
                     f"<td{tooltip}><span style='color:{ac};font-weight:600;font-size:11px'>{_resc(alert)}</span></td>"
                     f"<td style='font-size:11px;color:#64748b;max-width:300px'>{desc}</td></tr>")
        sections.append(
            f'<h2>{sn}. Portfel tokenów</h2>{_sdesc("portfel")}'
            f'<table class="data-table"><thead><tr>'
            f'<th>Token</th><th>Nazwa</th><th>Kategoria</th>'
            f'<th>Wpływy</th><th>Wypływy</th><th>Saldo netto</th><th>TX</th>'
            f'<th>Alert</th><th>Opis</th>'
            f'</tr></thead><tbody>{rows}</tbody></table>'
        )

    # ── Charts: Saldo w czasie + Graf przepływu ──
    charts = r.get("charts", {})
    graph = r.get("graph", {})
    balance_data = charts.get("balance_timeline", {})
    has_balance = bool(balance_data and balance_data.get("labels"))
    has_graph = bool(graph and graph.get("nodes"))
    if has_balance or has_graph:
        sn += 1
        chart_html = ""

        # Balance timeline chart (Chart.js) with log-normalization
        if has_balance:
            import json as _json
            chart_html += (
                '<h3>Saldo w czasie</h3>'
                '<div style="max-width:1060px;margin:0 auto 24px">'
                '<canvas id="report_balance_chart" width="1060" height="400"></canvas>'
                '</div>'
                '<script>'
                f'var _btData={_json.dumps(balance_data, ensure_ascii=False)};'
                'document.addEventListener("DOMContentLoaded",function(){'
                'var c=document.getElementById("report_balance_chart");if(!c)return;'
                'var colors=["#2563eb","#dc2626","#22c55e","#f59e0b","#8b5cf6","#ec4899","#14b8a6","#f97316","#06b6d4","#84cc16"];'
                # Normalization logic matching the program view
                'var dsets=_btData.datasets||[];'
                'var maxPer=dsets.map(function(d){return Math.max.apply(null,(d.data||[]).map(function(v){return Math.abs(v||0)}))|| 0.0001;});'
                'var gMax=Math.max.apply(null,maxPer);var gMin=Math.min.apply(null,maxPer);'
                'var needsNorm=gMax>0&&gMin>0&&(gMax/gMin)>50;'
                'var ds,yOpts;'
                'if(needsNorm){'
                'var logMax=maxPer.map(function(m){return Math.log10(m+1);});'
                'var logGMax=Math.max.apply(null,logMax);'
                'ds=dsets.map(function(d,i){'
                'var tMax=maxPer[i]||1;var ls=logMax[i]/(logGMax||1);var ceil=20+ls*80;'
                'return{label:d.token+" (skala)",data:(d.data||[]).map(function(v){return((v||0)/tMax)*ceil;}),'
                'borderColor:colors[i%colors.length],backgroundColor:"transparent",'
                'borderWidth:Math.max(1,Math.min(3,ls*3)),pointRadius:0,tension:0.3};'
                '});'
                'yOpts={beginAtZero:true,max:105,title:{display:true,text:"Skala relatywna (log)"},ticks:{callback:function(v){return Math.round(v)+"%";},font:{size:10}}};'
                '}else{'
                'ds=dsets.map(function(d,i){'
                'return{label:d.token,data:d.data,borderColor:colors[i%colors.length],backgroundColor:"transparent",borderWidth:1.5,pointRadius:0,tension:0.3};'
                '});'
                'yOpts={ticks:{font:{size:10}}};'
                '}'
                'new Chart(c,{type:"line",data:{labels:_btData.labels,datasets:ds},'
                'options:{responsive:true,plugins:{legend:{position:"bottom",labels:{font:{size:10}}}},'
                'scales:{x:{ticks:{maxTicksLimit:20,font:{size:9}}},y:yOpts}}});'
                '});'
                '</script>'
            )

        # Flow graph (Cytoscape.js - same rendering as program view)
        if has_graph:
            nodes = graph.get("nodes", [])
            edges = graph.get("edges", [])
            if nodes and edges:
                import json as _json2
                risk_cm = {"critical": "#dc2626", "high": "#f97316", "medium": "#eab308", "low": "#64748b"}
                cy_elements = []
                for node in nodes:
                    d = node.get("data", {})
                    nc = risk_cm.get(d.get("risk_level", "low"), "#64748b")
                    size = max(20, min(60, 20 + (d.get("tx_count", 0) or 0) * 2))
                    cy_elements.append({"group": "nodes", "data": {
                        "id": d["id"], "label": d.get("label", d["id"][:10]),
                        "color": nc, "size": size,
                    }})
                for edge in edges[:200]:
                    d = edge.get("data", {})
                    amt = d.get("amount", 0)
                    label = d.get("token", "")
                    if amt:
                        label += f" {amt:,.2f}"
                    ec = "#ef4444" if d.get("risk") else "#475569"
                    w = max(1, min(6, d.get("count", 1) or 1))
                    cy_elements.append({"group": "edges", "data": {
                        "source": d["source"], "target": d["target"],
                        "label": label, "width": w, "color": ec,
                    }})
                chart_html += (
                    '<h3>Graf przepływu transakcji</h3>'
                    '<div id="report_graph" style="width:100%;height:550px;background:#fafbfc;border:1px solid #e2e8f0;border-radius:8px"></div>'
                    '<script src="https://cdn.jsdelivr.net/npm/cytoscape@3.30.4/dist/cytoscape.min.js"></script>'
                    '<script>'
                    f'var _gElems={_json2.dumps(cy_elements, ensure_ascii=False)};'
                    'document.addEventListener("DOMContentLoaded",function(){'
                    'var c=document.getElementById("report_graph");if(!c)return;'
                    'cytoscape({container:c,elements:_gElems,'
                    'style:['
                    '{selector:"node",style:{"background-color":"data(color)","label":"data(label)","color":"#1e293b",'
                    '"font-size":"10px","text-valign":"bottom","text-margin-y":4,"width":"data(size)","height":"data(size)",'
                    '"border-width":1,"border-color":"#94a3b8"}},'
                    '{selector:"edge",style:{"line-color":"data(color)","target-arrow-color":"data(color)",'
                    '"target-arrow-shape":"triangle","curve-style":"bezier","width":"data(width)",'
                    '"label":"data(label)","font-size":"8px","color":"#64748b","text-rotation":"autorotate",'
                    '"text-margin-y":-8,"text-opacity":0.8}}'
                    '],'
                    'layout:{name:"cose",animate:false,nodeDimensionsIncludeLabels:true,'
                    'nodeRepulsion:function(){return 6000;},idealEdgeLength:function(){return 120;}},'
                    'userZoomingEnabled:false,userPanningEnabled:false'
                    '});'
                    '});'
                    '</script>'
                )

        sections.append(f"<h2>{sn}. Wykresy</h2>{_sdesc('wykresy')}{chart_html}")

    # ── next. Kontrahenci i transfery ──
    cps = bs.get("counterparties", {})
    pay_cps = fr.get("binance_pay_counterparties", {})
    phones = r.get("detected_phones", [])
    if cps or pay_cps or phones:
        sn += 1
        ct_html = ""
        if cps:
            rows = ""
            for uid, c in sorted(cps.items(), key=lambda x: x[1].get("tx_count", 0), reverse=True):
                period = f"{(c.get('first_seen', '') or '')[:10]} - {(c.get('last_seen', '') or '')[:10]}"
                rows += f"<tr><td><code>{_resc(uid)}</code></td><td>{c.get('tx_count', 0)}</td><td class='num'>{c.get('total_in', 0):.4f}</td><td class='num'>{c.get('total_out', 0):.4f}</td><td>{_resc(', '.join(c.get('tokens', [])))}</td><td>{_resc(', '.join(c.get('sources', [])))}</td><td style='font-size:11px'>{_resc(period)}</td></tr>"
            ct_html += f'<h3>Kontrahenci wewnętrzni Binance</h3><table class="data-table"><thead><tr><th>User ID</th><th>TX</th><th>Wpływy</th><th>Wypływy</th><th>Tokeny</th><th>Źródło</th><th>Okres</th></tr></thead><tbody>{rows}</tbody></table>'
            ct_html += f'<p>Transfery wewnętrzne: <b>{bs.get("internal_transfer_count", 0)}</b> | Depozyty zewnętrzne: <b>{bs.get("external_deposit_count", 0)}</b> | Wypłaty zewnętrzne: <b>{bs.get("external_withdrawal_count", 0)}</b></p>'

        if pay_cps:
            rows = ""
            for k, c in sorted(pay_cps.items(), key=lambda x: x[1].get("count", 0), reverse=True):
                rows += f"<tr><td><code>{_resc(k)}</code></td><td>{_resc(c.get('wallet_id', ''))}</td><td>{c.get('count', 0)}</td><td class='num'>{c.get('in', 0):.4f}</td><td class='num'>{c.get('out', 0):.4f}</td><td>{_resc(', '.join(c.get('tokens', [])))}</td></tr>"
            ct_html += f'<h3>Binance Pay (C2C)</h3><table class="data-table"><thead><tr><th>Binance ID</th><th>Wallet ID</th><th>TX</th><th>Wpływy</th><th>Wypływy</th><th>Tokeny</th></tr></thead><tbody>{rows}</tbody></table>'

        if phones:
            rows = ""
            for p in phones:
                ctx_text = ""
                if p.get("contexts"):
                    c0 = p["contexts"][0]
                    ctx_text = f"{c0.get('tx_type', '')} {c0.get('token', '')} {c0.get('timestamp', '')} (pole: {c0.get('field', '')})"
                rows += f"<tr><td style='font-family:monospace;font-weight:600'>{_resc(p['number'])}</td><td>{_resc(p.get('country_name', ''))}</td><td>{_resc(p.get('country_iso', ''))}</td><td>{p.get('occurrences', 0)}</td><td style='font-size:11px'>{_resc(ctx_text)}</td></tr>"
            ct_html += f'<h3>Zidentyfikowane numery telefonów</h3><table class="data-table"><thead><tr><th>Numer</th><th>Kraj</th><th>ISO</th><th>Wystąpienia</th><th>Kontekst</th></tr></thead><tbody>{rows}</tbody></table>'

        sections.append(f"<h2>{sn}. Kontrahenci i transfery</h2>{_sdesc('kontrahenci')}{ct_html}")

    # ── VII. Adresy on-chain ──
    ext_src = fr.get("external_source_addresses", [])
    ext_dst = fr.get("external_dest_addresses", [])
    wallets = r.get("wallets", [])
    dep_addrs = bs.get("deposit_addresses", [])
    if ext_src or ext_dst or wallets or dep_addrs:
        sn += 1
        addr_html = ""

        if dep_addrs:
            rows = "".join(f"<tr><td><code style='word-break:break-all'>{_resc(a['address'])}</code></td><td>{_resc(a.get('chain', ''))}</td><td>{_resc(', '.join(a.get('tokens', [])))}</td></tr>" for a in dep_addrs)
            addr_html += f'<h3>Adresy depozytowe (portfele użytkownika)</h3><table class="data-table"><thead><tr><th>Adres</th><th>Sieć</th><th>Tokeny</th></tr></thead><tbody>{rows}</tbody></table>'

        # Merge external source + dest addresses (deduplicate, case-insensitive for EVM)
        if ext_src or ext_dst:
            def _nk(addr: str) -> str:
                a = addr.strip()
                return a.lower() if a.startswith("0x") or a.startswith("0X") else a

            addr_merged: Dict[str, Dict[str, Any]] = {}
            for a in ext_src:
                key = _nk(a["address"])
                if key in addr_merged:
                    m = addr_merged[key]
                    m["dep_count"] += a["count"]
                    m["dep_total"] += a["total"]
                    m["tokens"].update(a.get("tokens", []))
                    m["networks"].update(a.get("networks", []))
                else:
                    addr_merged[key] = {
                        "display": a["address"],
                        "dep_count": a["count"], "dep_total": a["total"],
                        "wd_count": 0, "wd_total": 0.0,
                        "tokens": set(a.get("tokens", [])),
                        "networks": set(a.get("networks", [])),
                    }
            for a in ext_dst:
                key = _nk(a["address"])
                if key in addr_merged:
                    m = addr_merged[key]
                    m["wd_count"] += a["count"]
                    m["wd_total"] += a["total"]
                    m["tokens"].update(a.get("tokens", []))
                    m["networks"].update(a.get("networks", []))
                else:
                    addr_merged[key] = {
                        "display": a["address"],
                        "dep_count": 0, "dep_total": 0.0,
                        "wd_count": a["count"], "wd_total": a["total"],
                        "tokens": set(a.get("tokens", [])),
                        "networks": set(a.get("networks", [])),
                    }
            sorted_addrs = sorted(addr_merged.items(),
                                  key=lambda x: x[1]["dep_total"] + x[1]["wd_total"], reverse=True)
            rows = ""
            for _key, m in sorted_addrs:
                display_addr = m.get("display", _key)
                direction = "dep+wd" if m["dep_count"] > 0 and m["wd_count"] > 0 else ("dep" if m["dep_count"] > 0 else "wd")
                dir_label = {"dep+wd": "&#x1F4E5;&#x1F4E4;", "dep": "&#x1F4E5;", "wd": "&#x1F4E4;"}[direction]
                dc = m["dep_count"] or "-"
                dt = f"{m['dep_total']:.4f}" if m["dep_count"] else "-"
                wc = m["wd_count"] or "-"
                wt = f"{m['wd_total']:.4f}" if m["wd_count"] else "-"
                rows += f"<tr><td><code style='word-break:break-all'>{_resc(display_addr)}</code></td><td style='text-align:center'>{dir_label}</td><td>{dc}</td><td class='num'>{dt}</td><td>{wc}</td><td class='num'>{wt}</td><td>{_resc(', '.join(sorted(m['tokens'])))}</td><td>{_resc(', '.join(sorted(m['networks'])))}</td></tr>"
            both_count = sum(1 for _, m in sorted_addrs if m["dep_count"] > 0 and m["wd_count"] > 0)
            note = f" (w tym <b>{both_count}</b> dwukierunkowych)" if both_count else ""
            addr_html += f'<h3>Adresy zewnętrzne ({len(sorted_addrs)} unikalnych{note})</h3><table class="data-table"><thead><tr><th>Adres</th><th>Kier.</th><th>Dep.TX</th><th>Dep.suma</th><th>Wyp.TX</th><th>Wyp.suma</th><th>Tokeny</th><th>Sieci</th></tr></thead><tbody>{rows}</tbody></table>'

        if wallets:
            rows = "".join(f"<tr><td><code style='word-break:break-all'>{_resc(w['address'])}</code></td><td>{_resc(w.get('label', ''))}</td><td>{w.get('tx_count', 0)}</td><td class='num'>{w.get('total_received', 0):.4f}</td><td class='num'>{w.get('total_sent', 0):.4f}</td><td>{_resc(w.get('risk_level', ''))}</td></tr>" for w in wallets[:100])
            addr_html += f'<h3>Portfele</h3><table class="data-table"><thead><tr><th>Adres</th><th>Etykieta</th><th>TX</th><th>Otrzymane</th><th>Wysłane</th><th>Ryzyko</th></tr></thead><tbody>{rows}</tbody></table>'

        sections.append(f"<h2>{sn}. Adresy on-chain</h2>{_sdesc('adresy')}{addr_html}")

    # ── VIII. Analiza kryminalistyczna ──
    pts = fr.get("pass_through_detection", [])
    priv = fr.get("privacy_coin_usage", {})
    mining = fr.get("mining_patterns", [])
    margin = fr.get("margin_analysis", {})
    if pts or priv or mining or margin:
        sn += 1
        forens_html = ""

        if pts:
            pt_count = fr.get("pass_through_count", len(pts))
            rows = "".join(
                f"<tr><td>{_resc(p['deposit_time'][:16])}</td><td class='num'>{p['deposit_amount']:.4f}</td><td>{_resc(p['deposit_token'])}</td><td>{_resc(p.get('deposit_from', ''))}</td>"
                f"<td>{_resc(p['withdrawal_time'][:16])}</td><td class='num'>{p['withdrawal_amount']:.4f}</td><td>{_resc(p['withdrawal_token'])}</td><td>{_resc(p.get('withdrawal_to', ''))}</td>"
                f"<td>{p['delay_hours']}h</td></tr>"
                for p in pts
            )
            forens_html += f'<h3>Przeloty tranzytowe (pass-through)</h3><p>Wykryto {pt_count} potencjalnych przepływów (depozyt → wypłata w ciągu 24h).</p><table class="data-table"><thead><tr><th>Dep. czas</th><th>Dep. kwota</th><th>Token</th><th>Od</th><th>Wyp. czas</th><th>Wyp. kwota</th><th>Token</th><th>Do</th><th>Opóźn.</th></tr></thead><tbody>{rows}</tbody></table>'

        if priv:
            rows = ""
            for coin, p in priv.items():
                rows += f"<tr><td style='color:#f59e0b;font-weight:bold'>{_resc(coin)}</td><td>{p.get('deposits', 0)}</td><td class='num'>{p.get('deposit_amount', 0):.4f}</td><td>{p.get('withdrawals', 0)}</td><td class='num'>{p.get('withdrawal_amount', 0):.4f}</td><td>{p.get('trades', 0)}</td><td>{p.get('unique_source_addresses', 0)}</td></tr>"
            forens_html += f'<h3>Kryptowaluty prywatności</h3><table class="data-table"><thead><tr><th>Moneta</th><th>Dep.</th><th>Kwota dep.</th><th>Wyp.</th><th>Kwota wyp.</th><th>Transakcje</th><th>Unik. adresy</th></tr></thead><tbody>{rows}</tbody></table>'

        if mining:
            rows = "".join(f"<tr><td><code style='word-break:break-all'>{_resc(m['address'])}</code></td><td>{_resc(m['token'])}</td><td>{m['count']}</td><td class='num'>{m['total']:.8f}</td><td class='num'>{m['avg']:.8f}</td></tr>" for m in mining)
            forens_html += f'<h3>Wzorce górnicze</h3><table class="data-table"><thead><tr><th>Adres</th><th>Token</th><th>TX</th><th>Suma</th><th>Średnia</th></tr></thead><tbody>{rows}</tbody></table>'

        if margin and margin.get("total_orders"):
            forens_html += f'<h3>Handel z dźwignią (margin)</h3><table class="info-table">'
            for lbl, val in [("User IDs", ", ".join(margin.get("user_ids", []))),
                             ("Łącznie zleceń", margin.get("total_orders")),
                             ("Zrealizowane", margin.get("filled_orders")),
                             ("Anulowane", margin.get("cancelled_orders")),
                             ("Kupno", margin.get("buy_count")),
                             ("Sprzedaż", margin.get("sell_count"))]:
                if val:
                    forens_html += f"<tr><th>{_resc(lbl)}</th><td>{_resc(str(val))}</td></tr>"
            forens_html += "</table>"
            if margin.get("top_markets"):
                mk = "".join(f"<tr><td>{_resc(m)}</td><td>{c}</td></tr>" for m, c in margin["top_markets"].items())
                forens_html += f'<h4>Top rynki margin</h4><table class="data-table"><thead><tr><th>Rynek</th><th>Zlecenia</th></tr></thead><tbody>{mk}</tbody></table>'

        sections.append(f"<h2>{sn}. Analiza kryminalistyczna</h2>{_sdesc('kryminalistyczna')}{forens_html}")

    # ── IX. Bezpieczeństwo konta ──
    al = fr.get("access_log_analysis", {})
    devs = fr.get("device_fingerprints", [])
    cards = fr.get("card_info", [])
    card_tl = fr.get("card_geo_timeline", [])
    if al.get("total_entries") or devs or cards or card_tl:
        sn += 1
        sec_html = ""

        if al.get("total_entries"):
            sec_html += f'<h3>Logi dostępu</h3><p>Łącznie wpisów: <b>{al["total_entries"]}</b>, unikalne IP: <b>{al.get("unique_ips", 0)}</b>, okres: {_resc(al.get("first_login", "")[:10])} - {_resc(al.get("last_login", "")[:10])}</p>'
            if al.get("foreign_login_count"):
                sec_html += f'<p style="color:#dc2626"><b>Zagraniczne loginy: {al["foreign_login_count"]}</b> (poza {_resc(al.get("primary_country", "?"))})</p>'
            geos = al.get("geolocations", {})
            if geos:
                rows = "".join(f"<tr><td>{_resc(g)}</td><td>{c}</td></tr>" for g, c in geos.items())
                sec_html += f'<h4>Geolokalizacje</h4><table class="data-table"><thead><tr><th>Lokalizacja</th><th>Loginy</th></tr></thead><tbody>{rows}</tbody></table>'
            ips = al.get("top_ips", {})
            if ips:
                rows = "".join(f"<tr><td><code>{_resc(ip)}</code></td><td>{c}</td></tr>" for ip, c in ips.items())
                sec_html += f'<h4>Najczęstsze IP</h4><table class="data-table"><thead><tr><th>IP</th><th>Loginy</th></tr></thead><tbody>{rows}</tbody></table>'

        if devs:
            rows = "".join(f"<tr><td>{_resc(d.get('device', ''))}</td><td>{_resc(d.get('client', ''))}</td><td><code>{_resc(d.get('ip', ''))}</code></td><td>{_resc(d.get('geo', ''))}</td><td>{_resc((d.get('last_used', '') or '')[:10])}</td><td>{_resc(d.get('status', ''))}</td></tr>" for d in devs)
            sec_html += f'<h3>Zatwierdzone urządzenia</h3><table class="data-table"><thead><tr><th>Urządzenie</th><th>Klient</th><th>IP</th><th>Geo</th><th>Ostatnie użycie</th><th>Status</th></tr></thead><tbody>{rows}</tbody></table>'

        if cards:
            rows = "".join(f"<tr><td><code>{_resc(c.get('card_number', ''))}</code></td><td>{_resc(c.get('card_type', ''))}</td><td>{_resc(c.get('status', ''))}</td><td>{_resc((c.get('created', '') or '')[:10])}</td></tr>" for c in cards)
            sec_html += f'<h3>Karty Binance</h3><table class="data-table"><thead><tr><th>Numer</th><th>Typ</th><th>Status</th><th>Utworzono</th></tr></thead><tbody>{rows}</tbody></table>'

        card_spending = bs.get("card_spending", {})
        card_merchants = bs.get("card_merchants", {})
        if card_spending:
            sec_html += '<h4>Wydatki kartą</h4><p>' + ", ".join(f"{_resc(k)}: <b>{v:.2f}</b>" for k, v in card_spending.items()) + '</p>'
        if card_merchants:
            rows = "".join(f"<tr><td>{_resc(m)}</td><td class='num'>{v:.2f}</td></tr>" for m, v in sorted(card_merchants.items(), key=lambda x: x[1], reverse=True))
            sec_html += f'<h4>Merchants</h4><table class="data-table"><thead><tr><th>Merchant</th><th>Kwota</th></tr></thead><tbody>{rows}</tbody></table>'

        if card_tl:
            rows = "".join(f"<tr><td>{_resc(t.get('timestamp', '')[:16])}</td><td>{_resc(t.get('merchant', ''))}</td><td class='num'>{t.get('amount', 0):.2f}</td><td>{_resc(t.get('currency', ''))}</td><td>{_resc(t.get('status', ''))}</td></tr>" for t in card_tl)
            sec_html += f'<h3>Oś czasu transakcji kartą</h3><table class="data-table"><thead><tr><th>Data</th><th>Merchant</th><th>Kwota</th><th>Waluta</th><th>Status</th></tr></thead><tbody>{rows}</tbody></table>'

        sections.append(f"<h2>{sn}. Bezpieczeństwo konta</h2>{_sdesc('bezpieczenstwo')}{sec_html}")

    # ── X. Analiza czasowa ──
    ta = fr.get("temporal_analysis", {})
    if ta and ta.get("active_span_days"):
        sn += 1
        t_html = f'<table class="info-table">'
        t_html += f'<tr><th>Okres aktywności</th><td>{ta.get("active_span_days", 0)} dni</td></tr>'
        t_html += f'<tr><th>Aktywne dni</th><td>{ta.get("active_days", 0)} ({ta.get("activity_density", 0)}%)</td></tr>'
        t_html += f'<tr><th>Szczytowa godzina</th><td>{ta.get("peak_hour", "?")}:00 ({ta.get("peak_hour_count", 0)} tx)</td></tr>'
        t_html += f'<tr><th>Aktywność nocna (0-5)</th><td>{ta.get("night_activity_count", 0)} ({ta.get("night_activity_ratio", 0)}%)</td></tr>'
        t_html += f'<tr><th>Weekend / dni robocze</th><td>{ta.get("weekend_count", 0)} / {ta.get("weekday_count", 0)} ({ta.get("weekend_ratio", 0)}%)</td></tr>'
        t_html += '</table>'
        bursts = ta.get("burst_days", [])
        if bursts:
            rows = "".join(f"<tr><td>{_resc(b['date'])}</td><td class='num' style='color:#dc2626'>{b['tx_count']}</td></tr>" for b in bursts)
            t_html += f'<h3>Dni z nagłą aktywnością (&gt;50 tx)</h3><table class="data-table"><thead><tr><th>Data</th><th>TX</th></tr></thead><tbody>{rows}</tbody></table>'
        dorm = ta.get("dormancy_periods", [])
        if dorm:
            rows = "".join(f"<tr><td>{_resc(d['from'])}</td><td>{_resc(d['to'])}</td><td class='num'>{d['days']}</td></tr>" for d in dorm)
            t_html += f'<h3>Okresy uśpienia (&gt;7 dni)</h3><table class="data-table"><thead><tr><th>Od</th><th>Do</th><th>Dni</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Analiza czasowa</h2>{_sdesc('czasowa')}{t_html}")

    # ── X. Łańcuchy konwersji ──
    cc = fr.get("conversion_chains", {})
    if cc and cc.get("edges"):
        sn += 1
        c_html = f'<p><b>Unikalne pary konwersji:</b> {cc.get("unique_swap_pairs", 0)}'
        if cc.get("fiat_entry_tokens"):
            c_html += f' | <b>Fiat wejściowe:</b> {_resc(", ".join(cc["fiat_entry_tokens"]))}'
        c_html += '</p>'
        rows = "".join(f"<tr><td><b>{_resc(e['from'])}</b></td><td>→</td><td><b>{_resc(e['to'])}</b></td><td class='num'>{e['volume']:.4f}</td></tr>" for e in cc["edges"])
        c_html += f'<table class="data-table"><thead><tr><th>Z tokenu</th><th></th><th>Na token</th><th>Wolumen</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Łańcuchy konwersji tokenów</h2>{_sdesc('konwersje')}{c_html}")

    # ── X. Structuring ──
    sd = fr.get("structuring_detection", {})
    s_alerts = sd.get("alerts", [])
    s_freq = sd.get("frequent_amounts", [])
    if s_alerts or s_freq:
        sn += 1
        s_html = ''
        if s_alerts:
            s_html += f'<p style="color:#dc2626;font-weight:bold">⚠️ Wykryto {sd.get("alert_count", len(s_alerts))} potencjalnych przypadków structuringu</p>'
            rows = "".join(f"<tr><td>{_resc(a['date'])}</td><td>{_resc(a['type'])}</td><td>{a['threshold']}</td><td>{a['count']}</td><td>{', '.join(str(x) for x in a.get('amounts', []))}</td><td class='num'>{a['daily_total']:.2f}</td></tr>" for a in s_alerts)
            s_html += f'<table class="data-table"><thead><tr><th>Data</th><th>Typ</th><th>Próg</th><th>TX</th><th>Kwoty</th><th>Suma</th></tr></thead><tbody>{rows}</tbody></table>'
        if s_freq:
            rows = "".join(f"<tr><td class='num'>{f['amount']:.0f}</td><td class='num'>{f['count']}</td></tr>" for f in s_freq)
            s_html += f'<h3>Najczęściej używane kwoty</h3><table class="data-table"><thead><tr><th>Kwota</th><th>Wystąpienia</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Wykrywanie structuringu (smurfing)</h2>{_sdesc('structuring')}{s_html}")

    # ── X. Wash trading ──
    wt = fr.get("wash_trading", {})
    w_rev = wt.get("rapid_reversals", [])
    w_net = wt.get("zero_net_markets", [])
    if w_rev or w_net:
        sn += 1
        w_html = ''
        if w_net:
            rows = "".join(f"<tr><td><b>{_resc(m['market'])}</b></td><td class='num'>{m['gross_volume']:.4f}</td><td class='num' style='color:#dc2626'>{m['net_position']:.4f}</td><td class='num'>{m['net_ratio']}%</td><td class='num'>{m['buys']:.4f}</td><td class='num'>{m['sells']:.4f}</td></tr>" for m in w_net)
            w_html += f'<h3>Rynki z zerową pozycją netto</h3><table class="data-table"><thead><tr><th>Rynek</th><th>Wol. brutto</th><th>Poz. netto</th><th>Net%</th><th>Kupno</th><th>Sprzedaż</th></tr></thead><tbody>{rows}</tbody></table>'
        if w_rev:
            rows = "".join(f"<tr><td>{_resc(w['market'])}</td><td>{_resc(w['time1'][:16])}</td><td>{_resc(w['side1'])}</td><td class='num'>{w['amount1']}</td><td>{_resc(w['time2'][:16])}</td><td>{_resc(w['side2'])}</td><td class='num'>{w['amount2']}</td><td class='num'>{w['delay_sec']}</td></tr>" for w in w_rev[:50])
            w_html += f'<h3>Szybkie odwrócenia (&lt;5 min) - {wt.get("rapid_reversal_count", len(w_rev))} wykrytych</h3><table class="data-table"><thead><tr><th>Rynek</th><th>Czas 1</th><th>Strona</th><th>Kwota</th><th>Czas 2</th><th>Strona</th><th>Kwota</th><th>Opóźn.(s)</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Wash trading</h2>{_sdesc('wash')}{w_html}")

    # ── X. Fiat ramp ──
    fa = fr.get("fiat_ramp_analysis", {})
    if fa and (fa.get("fiat_deposit_count", 0) > 0 or fa.get("fiat_withdrawal_count", 0) > 0):
        sn += 1
        f_html = f'<table class="info-table">'
        f_html += f'<tr><th>Wpłaty fiat</th><td>{fa.get("fiat_deposit_count", 0)}</td></tr>'
        f_html += f'<tr><th>Wypłaty fiat</th><td>{fa.get("fiat_withdrawal_count", 0)}</td></tr>'
        f_html += f'<tr><th>Łącznie fiat IN</th><td>{fa.get("total_fiat_in", 0):.2f}</td></tr>'
        f_html += f'<tr><th>Łącznie fiat OUT</th><td>{fa.get("total_fiat_out", 0):.2f}</td></tr>'
        nf = fa.get("net_fiat_flow", 0)
        f_html += f'<tr><th>Saldo netto</th><td style="color:{"#22c55e" if nf >= 0 else "#dc2626"};font-weight:bold">{nf:.2f}</td></tr>'
        if fa.get("fiat_to_crypto_wd_hours") is not None:
            f_html += f'<tr><th>Fiat→crypto wypłata</th><td>{fa["fiat_to_crypto_wd_hours"]:.1f} godz.</td></tr>'
        f_html += '</table>'
        ci = fa.get("currencies_in", {})
        co = fa.get("currencies_out", {})
        if ci or co:
            all_c = sorted(set(list(ci.keys()) + list(co.keys())))
            rows = "".join(f"<tr><td><b>{_resc(c)}</b></td><td class='num'>{ci.get(c, 0):.2f}</td><td class='num'>{co.get(c, 0):.2f}</td></tr>" for c in all_c)
            f_html += f'<table class="data-table"><thead><tr><th>Waluta</th><th>Wpłaty</th><th>Wypłaty</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Analiza fiat on/off ramp</h2>{_sdesc('fiat')}{f_html}")

    # ── X. P2P ──
    p2p = fr.get("p2p_analysis", {})
    if p2p and p2p.get("total_count", 0) > 0:
        sn += 1
        p_html = f'<table class="info-table">'
        p_html += f'<tr><th>Transakcje P2P</th><td>{p2p["total_count"]}</td></tr>'
        p_html += f'<tr><th>% aktywności</th><td>{p2p.get("total_pct", 0)}%</td></tr>'
        p_html += f'<tr><th>Wolumen</th><td>{p2p.get("total_volume", 0):.2f}</td></tr>'
        p_html += f'<tr><th>Unikalni kontrahenci</th><td>{p2p.get("unique_counterparties", 0)}</td></tr>'
        p_html += '</table>'
        tops = p2p.get("top_counterparties", [])
        if tops:
            rows = "".join(f"<tr><td><code>{_resc(cp['id'])}</code></td><td class='num'>{cp['count']}</td><td class='num'>{cp['volume']:.4f}</td><td>{_resc(', '.join(cp.get('tokens', [])))}</td></tr>" for cp in tops)
            p_html += f'<h3>Top kontrahenci P2P</h3><table class="data-table"><thead><tr><th>ID</th><th>TX</th><th>Wolumen</th><th>Tokeny</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Analiza P2P</h2>{_sdesc('p2p')}{p_html}")

    # ── X. Velocity ──
    va = fr.get("velocity_analysis", {})
    if va and va.get("token_velocities"):
        sn += 1
        v_html = f'<table class="info-table">'
        v_html += f'<tr><th>Wpłaty / Wypłaty</th><td>{va.get("deposit_count", 0)} / {va.get("withdrawal_count", 0)}</td></tr>'
        v_html += f'<tr><th>Stosunek DEP/WD</th><td>{va.get("dep_wd_ratio", 0)}</td></tr>'
        v_html += f'<tr><th>Hot wallet</th><td>{"TAK ⚠️" if va.get("has_hot_wallet_behavior") else "NIE"}</td></tr>'
        v_html += '</table>'
        rows = "".join(f"<tr><td><b>{_resc(t['token'])}</b></td><td class='num'>{t['avg_hold_hours']}</td><td class='num'>{t['min_hold_hours']}</td><td class='num'>{t['deposit_count']}</td><td class='num'>{t['withdrawal_count']}</td></tr>" for t in va["token_velocities"])
        v_html += f'<table class="data-table"><thead><tr><th>Token</th><th>Śr. czas (godz.)</th><th>Min (godz.)</th><th>Wpłaty</th><th>Wypłaty</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Prędkość przepływu środków</h2>{_sdesc('velocity')}{v_html}")

    # ── X. Fee analysis ──
    fee = fr.get("fee_analysis", {})
    if fee and fee.get("fee_paying_tx_count", 0) > 0:
        sn += 1
        fe_html = f'<table class="info-table">'
        fe_html += f'<tr><th>TX z opłatami</th><td>{fee["fee_paying_tx_count"]}</td></tr>'
        fe_html += f'<tr><th>Opłaty w BNB</th><td>{fee.get("bnb_fee_count", 0)} ({fee.get("bnb_fee_ratio", 0)}%)</td></tr>'
        fe_html += '</table>'
        fees = fee.get("total_fees_by_token", {})
        if fees:
            rows = "".join(f"<tr><td><b>{_resc(tok)}</b></td><td class='num'>{val:.8f}</td></tr>" for tok, val in fees.items())
            fe_html += f'<table class="data-table"><thead><tr><th>Token opłaty</th><th>Suma opłat</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Analiza opłat (fees)</h2>{_sdesc('fees')}{fe_html}")

    # ── X. Network analysis ──
    na = fr.get("network_analysis", {})
    if na and na.get("networks"):
        sn += 1
        _HR = {"TRX", "TRON", "TRC20", "BSC", "BEP20", "BEP2"}
        n_html = f'<p><b>Unikalne sieci:</b> {na.get("unique_networks", 0)}'
        hr_nets = na.get("high_risk_networks", [])
        if hr_nets:
            n_html += f' | <span style="color:#dc2626"><b>Sieci wysokiego ryzyka:</b> {_resc(", ".join(n["network"] for n in hr_nets))}</span>'
        n_html += '</p>'
        rows = ""
        for n in na["networks"]:
            is_hr = n["network"].upper() in _HR
            style = 'background:#fef2f2' if is_hr else ''
            rows += f"<tr style='{style}'><td><b>{_resc(n['network'])}</b>{'⚠️' if is_hr else ''}</td><td class='num'>{n['deposits']}</td><td class='num'>{n['withdrawals']}</td><td class='num'>{n['total_tx']}</td><td class='num'>{n['dep_volume']:.4f}</td><td class='num'>{n['wd_volume']:.4f}</td></tr>"
        n_html += f'<table class="data-table"><thead><tr><th>Sieć</th><th>Wpłaty</th><th>Wypłaty</th><th>TX</th><th>Wol. wpłat</th><th>Wol. wypłat</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Analiza sieci blockchain</h2>{_sdesc('sieci')}{n_html}")

    # ── X. Extended security ──
    es = fr.get("extended_security", {})
    if es:
        sn += 1
        e_html = f'<table class="info-table">'
        e_html += f'<tr><th>Kraje logowań</th><td>{", ".join(es.get("login_countries", []))} ({es.get("login_country_count", 0)})</td></tr>'
        e_html += f'<tr><th>Podejrzane dni VPN</th><td>{es.get("vpn_suspect_days", 0)}</td></tr>'
        e_html += f'<tr><th>API Trading</th><td>{"Włączone" if es.get("api_trading_enabled") else "Wyłączone"}</td></tr>'
        e_html += f'<tr><th>Sub-konto</th><td>{"Tak" if es.get("has_sub_account") else "Nie"}</td></tr>'
        e_html += '</table>'
        vpn = es.get("vpn_suspects", [])
        if vpn:
            rows = "".join(f"<tr><td>{_resc(v['date'])}</td><td>{_resc(', '.join(v.get('countries', [])))}</td><td class='num' style='color:#dc2626'>{v['country_count']}</td><td class='num'>{v['login_count']}</td></tr>" for v in vpn)
            e_html += f'<h3>Podejrzenie VPN/proxy</h3><table class="data-table"><thead><tr><th>Data</th><th>Kraje</th><th>Ilość krajów</th><th>Loginy</th></tr></thead><tbody>{rows}</tbody></table>'
        sections.append(f"<h2>{sn}. Rozszerzona analiza bezpieczeństwa</h2>{_sdesc('ext_security')}{e_html}")

    # ── X. Transakcje ──
    txs = r.get("transactions", [])
    if txs:
        sn += 1
        rows = ""
        for tx in txs[:500]:
            rs = tx.get("risk_score", 0) or 0
            risk_style = f"color:#dc2626;font-weight:bold" if rs >= 50 else ""
            to_addr = tx.get("counterparty", "") or tx.get("to", "") or ""
            from_addr = tx.get("from", "") or ""
            rows += f"<tr><td>{_resc((tx.get('timestamp', '') or '')[:16])}</td><td>{_resc(tx.get('tx_type', ''))}</td><td>{_resc(tx.get('token', ''))}</td><td class='num'>{tx.get('amount', 0):.4f}</td><td style='word-break:break-all;font-size:10px'>{_resc(from_addr)}</td><td style='word-break:break-all;font-size:10px'>{_resc(to_addr)}</td><td style='{risk_style}'>{rs:.0f}</td><td>{_resc(', '.join(tx.get('risk_tags', [])))}</td></tr>"
        total = r.get("transactions_total", len(txs))
        note = f'<p class="muted">Wyświetlono {min(len(txs), 500)} z {total} transakcji.</p>' if total > 500 else ""
        sections.append(f'<h2>{sn}. Transakcje</h2>{_sdesc("transakcje")}<table class="data-table"><thead><tr><th>Data</th><th>Typ</th><th>Token</th><th>Kwota</th><th>Od</th><th>Do/Kontrahent</th><th>Ryzyko</th><th>Tagi</th></tr></thead><tbody>{rows}</tbody></table>{note}')

    # ── Stopka: Słownik typów transakcji ──
    used_types: set = set()
    for tx in r.get("transactions", []):
        tt = tx.get("tx_type", "")
        if tt:
            used_types.add(tt.lower())
        cat = tx.get("category", "")
        if cat:
            used_types.add(cat.lower())

    legend_rows = ""
    for tt in sorted(used_types):
        desc = _TX_TYPE_DESCRIPTIONS.get(tt)
        if desc:
            legend_rows += f"<tr><td><strong>{_resc(tt)}</strong></td><td>{_resc(desc)}</td></tr>"

    if legend_rows:
        sn += 1
        sections.append(
            f'<h2>{sn}. Słownik typów transakcji</h2>'
            f'<table class="data-table"><thead><tr><th style="width:180px">Typ</th>'
            f'<th>Opis</th></tr></thead><tbody>{legend_rows}</tbody></table>'
        )

    body = "\n".join(sections)

    return f"""<!DOCTYPE html>
<html lang="pl">
<head>
<meta charset="utf-8">
<title>Raport Crypto - {_resc(filename)}</title>
<script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.7/dist/chart.umd.min.js"></script>
<style>
body {{ font-family: 'Segoe UI', Tahoma, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; color: #1e293b; font-size: 14px; }}
h1 {{ border-bottom: 3px solid #2563eb; padding-bottom: 8px; color: #1e293b; }}
h2 {{ color: #2563eb; border-bottom: 1px solid #e2e8f0; padding-bottom: 4px; margin-top: 32px; page-break-after: avoid; }}
h3 {{ color: #475569; margin-top: 20px; page-break-after: avoid; }}
h4 {{ color: #64748b; margin-top: 14px; }}
.info-table {{ border-collapse: collapse; margin: 12px 0; }}
.info-table th {{ text-align: left; padding: 6px 16px 6px 0; color: #64748b; font-weight: 600; white-space: nowrap; }}
.info-table td {{ padding: 6px 0; }}
.data-table {{ border-collapse: collapse; width: 100%; margin: 8px 0; font-size: 12px; }}
.data-table th {{ background: #f1f5f9; padding: 6px 8px; text-align: left; border: 1px solid #e2e8f0; font-weight: 600; }}
.data-table td {{ padding: 5px 8px; border: 1px solid #e2e8f0; }}
.data-table .num {{ text-align: right; font-variant-numeric: tabular-nums; }}
.data-table tr:nth-child(even) {{ background: #f8fafc; }}
code {{ background: #f1f5f9; padding: 1px 4px; border-radius: 3px; font-size: 11px; word-break: break-all; }}
.muted {{ color: #94a3b8; font-size: 12px; }}
.section-desc {{ color: #475569; font-size: 12px; line-height: 1.6; margin: 4px 0 14px; padding: 8px 12px; background: #f8fafc; border-left: 3px solid #cbd5e1; border-radius: 0 4px 4px 0; }}
.profile-card {{ padding: 10px 14px; margin: 8px 0; background: #f8fafc; border-radius: 6px; }}
.profile-card .desc {{ color: #64748b; font-size: 12px; margin-top: 2px; }}
.profile-card ul {{ margin: 6px 0 0; padding-left: 18px; font-size: 12px; }}
.footer {{ margin-top: 40px; padding-top: 12px; border-top: 1px solid #e2e8f0; color: #94a3b8; font-size: 11px; text-align: center; }}
@media print {{ body {{ font-size: 11px; }} h1 {{ font-size: 18px; }} h2 {{ font-size: 15px; break-inside: avoid; }} table {{ break-inside: auto; }} tr {{ break-inside: avoid; }} }}
</style>
</head>
<body>
<h1>Raport analizy kryptowalutowej</h1>
{body}
<div class="footer">
Wygenerowano: {now} &middot; AISTATE Crypto Analysis Module &middot; Plik: {_resc(filename)}
</div>
</body>
</html>"""


def _build_crypto_report_txt(r: Dict[str, Any]) -> str:
    """Build a plain-text report from crypto analysis results."""
    from datetime import datetime

    lines = []
    lines.append("=" * 70)
    lines.append("RAPORT ANALIZY KRYPTOWALUTOWEJ")
    lines.append("=" * 70)
    lines.append(f"Data wygenerowania: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")

    # Identification
    fr = r.get("forensic_report", {}) or {}
    ai = fr.get("account_info", {}) or {}
    em = r.get("exchange_meta", {}) or {}
    bs = r.get("binance_summary", {}) or {}

    meta = r.get("metadata", {}) or {}
    lines.append("--- IDENTYFIKACJA ---")
    for label, val in [
        ("Właściciel", ai.get("holder_name") or meta.get("account_holder")),
        ("Imię", ai.get("first_name")),
        ("Nazwisko", ai.get("last_name")),
        ("User ID", ai.get("user_id")),
        ("Email", ai.get("email")),
        ("Telefon", ai.get("phone")),
        ("Data urodzenia", ai.get("date_of_birth")),
        ("Płeć", ai.get("gender")),
        ("Kraj", ai.get("country") or meta.get("country")),
        ("Narodowość", ai.get("nationality")),
        ("Adres", ai.get("physical_address") or meta.get("street")),
        ("Miasto", ai.get("city") or meta.get("city")),
        ("Województwo/Stan", ai.get("state")),
        ("Kod pocztowy", ai.get("zip_code") or meta.get("postal_code")),
        ("KYC Level", ai.get("kyc_level")),
        ("VIP Level", ai.get("vip_level")),
        ("Data rejestracji", ai.get("registration_date")),
        ("Status konta", ai.get("account_status")),
        ("Typ dokumentu", ai.get("id_type")),
        ("Nr dokumentu", ai.get("id_number")),
        ("ID polecającego", ai.get("referral_id")),
        ("Platforma", em.get("exchange_name", r.get("source", ""))),
        ("Plik", r.get("filename")),
    ]:
        if val:
            lines.append(f"  {label}: {val}")
    lines.append("")

    # Summary
    lines.append("--- PODSUMOWANIE ---")
    lines.append(f"  Okres: {(r.get('date_from', '') or '')[:10]} - {(r.get('date_to', '') or '')[:10]}")
    lines.append(f"  Transakcje: {r.get('tx_count', 0)}")
    lines.append(f"  Ryzyko AML: {r.get('risk_score', 0):.1f}/100")
    lines.append("")

    # Behavior
    bp = r.get("behavior_profile", {})
    if bp and bp.get("profiles"):
        lines.append("--- PROFIL ZACHOWANIA ---")
        for p in bp["profiles"][:3]:
            if p["score"] < 15:
                continue
            lines.append(f"  {p.get('icon', '')} {p['label']} - {p['score']}%")
            for rr in p.get("reasons", []):
                lines.append(f"    - {rr}")
        lines.append("")

    # Risk
    risk_reasons = r.get("risk_reasons", [])
    if risk_reasons:
        lines.append("--- CZYNNIKI RYZYKA ---")
        for rr in risk_reasons:
            lines.append(f"  - {rr}")
        lines.append("")

    # Tokens
    tokens = r.get("tokens", {})
    if tokens:
        lines.append("--- PORTFEL TOKENÓW ---")
        tc_txt = r.get("token_classification", {}) or {}
        lines.append(f"  {'Token':<8} {'Nazwa':<16} {'Rank':>5} {'Kategoria':<22} {'Wpływy':>14} {'Wypływy':>14} {'Saldo':>14} {'TX':>6}  {'Alert':<8} Opis")
        for tok, s in sorted(tokens.items(), key=lambda x: x[1].get("count", 0), reverse=True):
            net = (s.get("received", 0) or 0) - (s.get("sent", 0) or 0)
            info = tc_txt.get(tok, {})
            name = info.get("name", "")[:15]
            rank = f"#{info['rank']}" if info.get("rank") else "-"
            cat = info.get("category", "")[:20]
            alert = info.get("alert_level", "NORMAL")
            desc = info.get("description", "")[:60]
            lines.append(f"  {tok:<8} {name:<16} {rank:>5} {cat:<22} {s.get('received', 0):>14.4f} {s.get('sent', 0):>14.4f} {net:>14.4f} {s.get('count', 0):>6}  {alert:<8} {desc}")
        lines.append("")

    # Phone numbers
    phones = r.get("detected_phones", [])
    if phones:
        lines.append("--- ZIDENTYFIKOWANE NUMERY TELEFONÓW ---")
        for p in phones:
            lines.append(f"  {p['number']:<18} {p.get('country_name', '?'):<20} ({p.get('country_iso', '?')}) x{p.get('occurrences', 0)}")
        lines.append("")

    # Counterparties
    cps = bs.get("counterparties", {})
    if cps:
        lines.append("--- KONTRAHENCI ---")
        for uid, c in sorted(cps.items(), key=lambda x: x[1].get("tx_count", 0), reverse=True):
            lines.append(f"  UID: {uid}  TX: {c.get('tx_count', 0)}  IN: {c.get('total_in', 0):.4f}  OUT: {c.get('total_out', 0):.4f}  Tokeny: {', '.join(c.get('tokens', []))}")
        lines.append("")

    # Addresses - merged, deduplicated (case-insensitive for EVM)
    ext_src = fr.get("external_source_addresses", [])
    ext_dst = fr.get("external_dest_addresses", [])
    if ext_src or ext_dst:
        def _nk_txt(addr: str) -> str:
            a = addr.strip()
            return a.lower() if a.startswith("0x") or a.startswith("0X") else a

        addr_m: Dict[str, Dict[str, Any]] = {}
        for a in ext_src:
            key = _nk_txt(a["address"])
            if key in addr_m:
                addr_m[key]["dc"] += a["count"]
                addr_m[key]["dt"] += a["total"]
                addr_m[key]["tok"].update(a.get("tokens", []))
            else:
                addr_m[key] = {"disp": a["address"], "dc": a["count"], "dt": a["total"], "wc": 0, "wt": 0.0, "tok": set(a.get("tokens", []))}
        for a in ext_dst:
            key = _nk_txt(a["address"])
            if key in addr_m:
                addr_m[key]["wc"] += a["count"]
                addr_m[key]["wt"] += a["total"]
                addr_m[key]["tok"].update(a.get("tokens", []))
            else:
                addr_m[key] = {"disp": a["address"], "dc": 0, "dt": 0.0, "wc": a["count"], "wt": a["total"], "tok": set(a.get("tokens", []))}
        lines.append("--- ADRESY ZEWNĘTRZNE (zjednoczone) ---")
        for _key, m in sorted(addr_m.items(), key=lambda x: x[1]["dt"] + x[1]["wt"], reverse=True):
            d = "DEP+WD" if m["dc"] > 0 and m["wc"] > 0 else ("DEP" if m["dc"] > 0 else "WD")
            lines.append(f"  {m['disp']}")
            lines.append(f"    Kier: {d}  Dep TX: {m['dc']}  Dep suma: {m['dt']:.4f}  Wyp TX: {m['wc']}  Wyp suma: {m['wt']:.4f}  Tokeny: {', '.join(sorted(m['tok']))}")
        lines.append("")

    # Temporal analysis
    ta = fr.get("temporal_analysis", {})
    if ta and ta.get("active_span_days"):
        lines.append("--- ANALIZA CZASOWA ---")
        lines.append(f"  Okres aktywności: {ta.get('active_span_days', 0)} dni")
        lines.append(f"  Aktywne dni: {ta.get('active_days', 0)} ({ta.get('activity_density', 0)}%)")
        lines.append(f"  Szczytowa godzina: {ta.get('peak_hour', '?')}:00 ({ta.get('peak_hour_count', 0)} tx)")
        lines.append(f"  Nocna aktywność (0-5): {ta.get('night_activity_count', 0)} ({ta.get('night_activity_ratio', 0)}%)")
        lines.append(f"  Weekend/robocze: {ta.get('weekend_count', 0)}/{ta.get('weekday_count', 0)} ({ta.get('weekend_ratio', 0)}%)")
        for b in ta.get("burst_days", []):
            lines.append(f"  BURST: {b['date']} - {b['tx_count']} tx")
        for d in ta.get("dormancy_periods", []):
            lines.append(f"  UŚPIENIE: {d['from']} → {d['to']} ({d['days']} dni)")
        lines.append("")

    # Conversion chains
    cc = fr.get("conversion_chains", {})
    if cc and cc.get("edges"):
        lines.append("--- ŁAŃCUCHY KONWERSJI ---")
        lines.append(f"  Unikalne pary: {cc.get('unique_swap_pairs', 0)}")
        for e in cc["edges"]:
            lines.append(f"  {e['from']:<8} → {e['to']:<8} vol: {e['volume']:.4f}")
        lines.append("")

    # Structuring
    sd = fr.get("structuring_detection", {})
    if sd.get("alerts"):
        lines.append("--- STRUCTURING / SMURFING ---")
        for a in sd["alerts"]:
            lines.append(f"  {a['date']} {a['type']:<12} próg:{a['threshold']}  x{a['count']}  suma:{a['daily_total']:.2f}")
        lines.append("")

    # Wash trading
    wt = fr.get("wash_trading", {})
    if wt.get("zero_net_markets") or wt.get("rapid_reversals"):
        lines.append("--- WASH TRADING ---")
        for m in wt.get("zero_net_markets", []):
            lines.append(f"  {m['market']:<16} brutto:{m['gross_volume']:.4f}  netto:{m['net_position']:.4f}  ({m['net_ratio']}%)")
        lines.append(f"  Szybkie odwrócenia: {wt.get('rapid_reversal_count', 0)}")
        lines.append("")

    # Fiat ramp
    fa_txt = fr.get("fiat_ramp_analysis", {})
    if fa_txt and (fa_txt.get("fiat_deposit_count", 0) > 0 or fa_txt.get("fiat_withdrawal_count", 0) > 0):
        lines.append("--- FIAT ON/OFF RAMP ---")
        lines.append(f"  Wpłaty fiat: {fa_txt.get('fiat_deposit_count', 0)}  Suma: {fa_txt.get('total_fiat_in', 0):.2f}")
        lines.append(f"  Wypłaty fiat: {fa_txt.get('fiat_withdrawal_count', 0)}  Suma: {fa_txt.get('total_fiat_out', 0):.2f}")
        lines.append(f"  Saldo netto: {fa_txt.get('net_fiat_flow', 0):.2f}")
        if fa_txt.get("fiat_to_crypto_wd_hours") is not None:
            lines.append(f"  Fiat→crypto wypłata: {fa_txt['fiat_to_crypto_wd_hours']:.1f} godz.")
        lines.append("")

    # P2P
    p2p_txt = fr.get("p2p_analysis", {})
    if p2p_txt and p2p_txt.get("total_count", 0) > 0:
        lines.append("--- ANALIZA P2P ---")
        lines.append(f"  Transakcje P2P: {p2p_txt['total_count']} ({p2p_txt.get('total_pct', 0)}% aktywności)")
        lines.append(f"  Wolumen: {p2p_txt.get('total_volume', 0):.2f}")
        lines.append(f"  Kontrahenci: {p2p_txt.get('unique_counterparties', 0)}")
        lines.append("")

    # Velocity
    va_txt = fr.get("velocity_analysis", {})
    if va_txt and va_txt.get("token_velocities"):
        lines.append("--- PRĘDKOŚĆ PRZEPŁYWU ---")
        lines.append(f"  DEP/WD: {va_txt.get('deposit_count', 0)}/{va_txt.get('withdrawal_count', 0)} (ratio: {va_txt.get('dep_wd_ratio', 0)})")
        lines.append(f"  Hot wallet: {'TAK' if va_txt.get('has_hot_wallet_behavior') else 'NIE'}")
        for t in va_txt["token_velocities"][:15]:
            lines.append(f"  {t['token']:<8} śr: {t['avg_hold_hours']} godz.  min: {t['min_hold_hours']} godz.")
        lines.append("")

    # Fee analysis
    fee_txt = fr.get("fee_analysis", {})
    if fee_txt and fee_txt.get("fee_paying_tx_count", 0) > 0:
        lines.append("--- ANALIZA OPŁAT ---")
        lines.append(f"  TX z opłatami: {fee_txt['fee_paying_tx_count']}")
        lines.append(f"  BNB fees: {fee_txt.get('bnb_fee_count', 0)} ({fee_txt.get('bnb_fee_ratio', 0)}%)")
        for tok, val in fee_txt.get("total_fees_by_token", {}).items():
            lines.append(f"  {tok:<8} {val:.8f}")
        lines.append("")

    # Network analysis
    na_txt = fr.get("network_analysis", {})
    if na_txt and na_txt.get("networks"):
        lines.append("--- ANALIZA SIECI BLOCKCHAIN ---")
        for n in na_txt["networks"]:
            lines.append(f"  {n['network']:<10} dep:{n['deposits']}  wd:{n['withdrawals']}  vol_dep:{n['dep_volume']:.4f}  vol_wd:{n['wd_volume']:.4f}")
        lines.append("")

    # Extended security
    es_txt = fr.get("extended_security", {})
    if es_txt:
        lines.append("--- ROZSZERZONE BEZPIECZEŃSTWO ---")
        lines.append(f"  Kraje logowań: {', '.join(es_txt.get('login_countries', []))}")
        lines.append(f"  Podejrzane dni VPN: {es_txt.get('vpn_suspect_days', 0)}")
        lines.append(f"  API Trading: {'Tak' if es_txt.get('api_trading_enabled') else 'Nie'}")
        lines.append(f"  Sub-konto: {'Tak' if es_txt.get('has_sub_account') else 'Nie'}")
        lines.append("")

    # Transactions
    txs = r.get("transactions", [])
    if txs:
        lines.append(f"--- TRANSAKCJE ({len(txs)}) ---")
        for tx in txs:
            ts = (tx.get("timestamp", "") or "")[:16]
            lines.append(f"  {ts}  {tx.get('tx_type', ''):<12} {tx.get('token', ''):<6} {tx.get('amount', 0):>14.4f}  {tx.get('counterparty', '') or tx.get('to', '') or ''}")
        lines.append("")

    # Słownik typów transakcji
    used_types: set = set()
    for tx in r.get("transactions", []):
        tt = tx.get("tx_type", "")
        if tt:
            used_types.add(tt.lower())
        cat = tx.get("category", "")
        if cat:
            used_types.add(cat.lower())
    legend = [(tt, _TX_TYPE_DESCRIPTIONS[tt]) for tt in sorted(used_types) if tt in _TX_TYPE_DESCRIPTIONS]
    if legend:
        lines.append("-" * 70)
        lines.append("SŁOWNIK TYPÓW TRANSAKCJI")
        lines.append("-" * 70)
        for tt, desc in legend:
            lines.append(f"  {tt:<20} - {desc}")
        lines.append("")

    lines.append("=" * 70)
    lines.append("Wygenerowano przez AISTATE Crypto Analysis Module")
    lines.append("=" * 70)

    return "\n".join(lines)


def _build_crypto_report_docx(r: Dict[str, Any]) -> bytes:
    """Build a DOCX report from crypto analysis results.

    Uses python-docx to create a formatted Word document with the same
    structure as the TXT report but with proper headings, tables and styles.
    """
    from io import BytesIO
    from datetime import datetime
    from docx import Document  # type: ignore
    from docx.shared import Pt, Cm, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    doc = Document()

    # Base styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    fr = r.get("forensic_report", {}) or {}
    ai = fr.get("account_info", {}) or {}
    em = r.get("exchange_meta", {}) or {}
    bs = r.get("binance_summary", {}) or {}
    risk_score = r.get("risk_score", 0)

    doc.add_heading("Raport analizy kryptowalutowej", level=0)
    p = doc.add_paragraph(f"Wygenerowano: {now}")
    p.runs[0].italic = True

    # ── Helper ──
    def add_kv_table(pairs):
        """Add a 2-column key-value table."""
        filtered = [(k, v) for k, v in pairs if v and str(v).strip()]
        if not filtered:
            return
        tbl = doc.add_table(rows=len(filtered), cols=2)
        tbl.style = "Table Grid"
        for i, (k, v) in enumerate(filtered):
            tbl.cell(i, 0).text = str(k)
            tbl.cell(i, 1).text = str(v)
            for cell in (tbl.cell(i, 0), tbl.cell(i, 1)):
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(1)
                    for run in para.runs:
                        run.font.size = Pt(10)
            # Bold label
            for run in tbl.cell(i, 0).paragraphs[0].runs:
                run.bold = True

    def add_data_table(headers, rows):
        """Add a table with header row and data rows."""
        if not rows:
            return
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        # Header
        for ci, h in enumerate(headers):
            tbl.cell(0, ci).text = h
            for run in tbl.cell(0, ci).paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        # Data
        for ri, row in enumerate(rows, 1):
            for ci, val in enumerate(row):
                tbl.cell(ri, ci).text = str(val) if val else ""
                for para in tbl.cell(ri, ci).paragraphs:
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(9)

    # Section descriptions (same dict as HTML report, defined in _build_crypto_report_html)
    _DOCX_DESC = {
        "podsumowanie": (
            "Podstawowe statystyki analizowanego konta kryptowalutowego. Sekcja zawiera informacje o łącznej liczbie transakcji, "
            "zakresie dat aktywności, listę obsługiwanych tokenów (kryptowalut) oraz źródło danych (np. nazwa giełdy, eksport z blockchainu). "
            "Te dane pozwalają szybko ocenić skalę aktywności - czy mamy do czynienia z kontem o kilku transakcjach, "
            "czy z intensywnie używanym rachunkiem z tysiącami operacji rozłożonych na lata."
        ),
        "profil": (
            "Na podstawie wzorców transakcji system automatycznie rozpoznaje typ użytkownika - np. inwestor długoterminowy (kupuje i trzyma), "
            "aktywny trader (częste kupna/sprzedaże), arbitrażysta (wykorzystuje różnice cenowe między giełdami) czy podmiot tranzytowy "
            "(środki szybko przechodzą przez konto). Profil jest wyznaczany algorytmicznie na podstawie częstotliwości transakcji, "
            "czasu trzymania tokenów, rodzajów operacji i ich wolumenów. Wynik zawiera procentowy poziom pewności oraz listę powodów, "
            "dla których dany profil został przypisany. Jeden użytkownik może wykazywać cechy kilku profili jednocześnie."
        ),
        "ryzyko": (
            "Kluczowa sekcja z punktu widzenia przeciwdziałania praniu pieniędzy (AML - Anti-Money Laundering). "
            "System analizuje wszystkie transakcje i przypisuje zagregowany wynik ryzyka w skali 0–100 punktów, "
            "gdzie 0 oznacza brak podejrzanych zachowań, a 100 - bardzo wysokie ryzyko. Ocena jest budowana z wielu czynników, m.in.: "
            "structuring (dzielenie kwot na mniejsze, aby ominąć progi raportowe), privacy coins (użycie kryptowalut utrudniających śledzenie), "
            "szybkie przeloty (środki wpływające i wypływające w ciągu godzin), kontrahenci wysokiego ryzyka."
        ),
        "portfel": (
            "Zestawienie wszystkich kryptowalut (tokenów), które pojawiły się na analizowanym koncie. "
            "Wpływy i wypływy to łączna wartość otrzymana i wysłana w danym tokenie; saldo netto to różnica. "
            "Kategoria klasyfikuje token (np. L1/infrastructure, stablecoin, payments/transfers, memecoin, privacy coin). "
            "Alert wskazuje poziom ryzyka tokena (NORMAL, MEDIUM, HIGH, CRITICAL); "
            "token oznaczony jako HIGH/CRITICAL wymaga szczególnej uwagi analityka."
        ),
        "wykresy": (
            "Graficzna prezentacja danych ułatwiająca szybkie wychwycenie trendów i anomalii. "
            "Wykres salda w czasie pokazuje jak zmieniało się saldo każdego tokena. "
            "Graf przepływu transakcji przedstawia powiązania między adresami i kontrahentami."
        ),
        "telefony": (
            "Numery telefonów wykryte w danych transakcyjnych, zidentyfikowane wraz z krajem pochodzenia. "
            "Powiązanie numerów telefonów z konkretnymi transakcjami może pomóc w identyfikacji kontrahentów."
        ),
        "kontrahenci": (
            "Lista kontrahentów wewnętrznych giełdy z podsumowaniem liczby transakcji, wolumenów wpływów i wypływów oraz używanych tokenów. "
            "Pozwala zidentyfikować najaktywniejszych partnerów handlowych i wykryć nietypowe powiązania."
        ),
        "adresy": (
            "Zestawienie adresów blockchain powiązanych z analizowanym kontem. "
            "Adresy depozytowe to portfele użytkownika na giełdzie. Adresy zewnętrzne to portfele spoza giełdy. "
            "Adresy dwukierunkowe mogą wskazywać na własne portfele użytkownika poza giełdą lub na bliskie relacje handlowe."
        ),
        "czasowa": (
            "Rozkład transakcji w podziale na godziny dnia i dni tygodnia. "
            "Pozwala wykryć aktywność automatyczną (boty), koordynację (wąskie okna czasowe) "
            "i nietypowe wzorce (np. wyłącznie nocna aktywność sugerująca inną strefę czasową)."
        ),
        "konwersje": (
            "Sekwencje szybkich zamian jednego tokena na drugi, np. PLN → BTC → XMR → USDT → EUR. "
            "Wieloetapowe konwersje w krótkim czasie to klasyczna technika layeringu - "
            "tworzenia warstw transakcji utrudniających prześledzenie pochodzenia środków."
        ),
        "structuring": (
            "Structuring (smurfing) polega na dzieleniu dużej kwoty na wiele mniejszych transakcji - "
            "często tuż poniżej progów raportowych (np. 15 000 EUR w UE, 10 000 USD w USA). "
            "Celem jest uniknięcie automatycznego zgłoszenia transakcji do jednostki analityki finansowej."
        ),
        "wash": (
            "Wash trading to transakcje, w których ten sam podmiot występuje jednocześnie jako kupujący i sprzedający. "
            "Celem jest sztuczne generowanie wolumenu obrotu - np. aby token wyglądał na bardziej płynny niż jest w rzeczywistości."
        ),
        "fiat": (
            "Punkty styku między tradycyjnym systemem finansowym a światem kryptowalut. "
            "On-ramp (wpłata fiat) to wpłaty z konta bankowego na giełdę; off-ramp to wypłaty z giełdy na konto bankowe. "
            "Asymetria wpłat i wypłat fiat może wskazywać na pranie pieniędzy lub transfer wartości."
        ),
        "p2p": (
            "Transakcje peer-to-peer to bezpośrednia wymiana kryptowalut między użytkownikami, "
            "często z pominięciem standardowej książki zleceń giełdy. Handel P2P bywa wykorzystywany "
            "do omijania procedur KYC (Know Your Customer) i AML."
        ),
        "velocity": (
            "Analiza szybkości przepływu środków przez konto. "
            "Średni czas trzymania to ile czasu token pozostaje na koncie; konta tranzytowe trzymają minuty lub godziny. "
            "Wskaźniki hot wallet oznaczają tokeny ze średnim trzymaniem poniżej 1 godziny."
        ),
        "fees": (
            "Zestawienie opłat transakcyjnych w podziale na tokeny. "
            "Pozwala oszacować koszty aktywności na koncie i wykryć transakcje z nietypowo wysokimi opłatami "
            "(celowe podnoszenie opłat może wskazywać na front-running)."
        ),
        "sieci": (
            "Statystyki wykorzystania różnych sieci blockchain. Pozwala określić preferencje sieciowe, "
            "wykryć bridge'owanie (przenoszenie środków między łańcuchami) "
            "oraz ocenić, czy użytkownik korzysta z sieci o niższym poziomie monitoringu AML."
        ),
        "ext_security": (
            "Pogłębiona analiza techniczna sesji i urządzeń: anomalie logowań, wykryte VPN/proxy, "
            "geolokalizacja kart - rozbieżności lokalizacji transakcji kartowych z logowaniami "
            "mogą wskazywać na kradzież tożsamości lub cloning karty."
        ),
        "transakcje": (
            "Kompletna lista wszystkich transakcji - baza dowodowa dla całej analizy. "
            "Każdy wiersz zawiera datę, typ operacji, token i kwotę, kontrahentów, "
            "poziom ryzyka i tagi klasyfikacyjne. Pozwala analitykowi na ręczną weryfikację "
            "dowolnej transakcji wychwycone przez algorytmy w sekcjach powyżej."
        ),
    }

    def _add_desc(key: str) -> None:
        """Add an italic description paragraph to the DOCX document."""
        txt = _DOCX_DESC.get(key, "")
        if txt:
            p = doc.add_paragraph(txt)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── 1. Identyfikacja ──
    meta = r.get("metadata", {}) or {}
    doc.add_heading("1. Identyfikacja podmiotu", level=1)
    add_kv_table([
        ("Właściciel konta", ai.get("holder_name") or meta.get("account_holder")),
        ("Imię", ai.get("first_name")),
        ("Nazwisko", ai.get("last_name")),
        ("User ID", ai.get("user_id")),
        ("Email", ai.get("email")),
        ("Telefon", ai.get("phone")),
        ("Data urodzenia", ai.get("date_of_birth")),
        ("Kraj", ai.get("country") or meta.get("country")),
        ("Narodowość", ai.get("nationality")),
        ("Adres", ai.get("physical_address") or meta.get("street")),
        ("Miasto", ai.get("city") or meta.get("city")),
        ("Kod pocztowy", ai.get("zip_code") or meta.get("postal_code")),
        ("KYC Level", ai.get("kyc_level")),
        ("VIP Level", ai.get("vip_level")),
        ("Data rejestracji", ai.get("registration_date")),
        ("Status konta", ai.get("account_status")),
        ("Typ dokumentu", ai.get("id_type")),
        ("Nr dokumentu", ai.get("id_number")),
        ("Platforma", em.get("exchange_name", r.get("source", ""))),
        ("Plik", r.get("filename")),
    ])

    # ── 2. Podsumowanie ──
    doc.add_heading("2. Podsumowanie ogólne", level=1)
    _add_desc("podsumowanie")
    date_from = (r.get("date_from", "") or "")[:10]
    date_to = (r.get("date_to", "") or "")[:10]
    add_kv_table([
        ("Okres analizy", f"{date_from} - {date_to}"),
        ("Transakcje", r.get("tx_count", 0)),
        ("Portfele/adresy", r.get("wallet_count", 0)),
        ("Kontrahenci", r.get("counterparty_count", 0)),
        ("Unikalne tokeny", len(r.get("tokens", {}))),
        ("Ryzyko AML", f"{risk_score:.1f}/100"),
    ])

    # ── 3. Profil zachowania ──
    bp = r.get("behavior_profile", {})
    if bp and bp.get("profiles"):
        doc.add_heading("3. Profil zachowania użytkownika", level=1)
        _add_desc("profil")
        for p in bp["profiles"][:5]:
            if p["score"] < 15:
                continue
            para = doc.add_paragraph()
            run = para.add_run(f"{p.get('icon', '')} {p['label']} - {p['score']}%")
            run.bold = True
            for reason in p.get("reasons", []):
                doc.add_paragraph(f"  • {reason}", style="List Bullet")

    # ── 4. Ryzyko AML ──
    risk_reasons = r.get("risk_reasons", [])
    if risk_reasons:
        doc.add_heading("4. Czynniki ryzyka AML", level=1)
        _add_desc("ryzyko")
        for rr in risk_reasons:
            doc.add_paragraph(f"• {rr}", style="List Bullet")

    # ── 5. Portfel tokenów ──
    tokens = r.get("tokens", {})
    tc = r.get("token_classification", {}) or {}
    if tokens:
        doc.add_heading("5. Portfel tokenów", level=1)
        _add_desc("portfel")
        t_rows = []
        for tok, s in sorted(tokens.items(), key=lambda x: x[1].get("count", 0), reverse=True):
            net = (s.get("received", 0) or 0) - (s.get("sent", 0) or 0)
            info = tc.get(tok, {})
            name = info.get("name", "")
            cat = info.get("category", "")
            alert = info.get("alert_level", "NORMAL")
            desc = info.get("description", "")
            t_rows.append([tok, name, cat,
                           f"{s.get('received', 0):.4f}", f"{s.get('sent', 0):.4f}",
                           f"{net:.4f}", str(s.get("count", 0)),
                           alert, desc])
        add_data_table(["Token", "Nazwa", "Kategoria",
                        "Wpływy", "Wypływy", "Saldo", "TX",
                        "Alert", "Opis"], t_rows)

    # ── Charts: Saldo w czasie + Graf przepływu ──
    charts = r.get("charts", {})
    graph = r.get("graph", {})
    balance_data = charts.get("balance_timeline", {})
    has_balance = bool(balance_data and balance_data.get("labels"))
    has_graph = bool(graph and graph.get("nodes") and graph.get("edges"))
    if has_balance or has_graph:
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            import matplotlib.dates as mdates

            doc.add_heading("6. Wykresy", level=1)
            _add_desc("wykresy")

            if has_balance:
                import math as _math
                fig, ax = plt.subplots(figsize=(8, 3.5))
                labels = balance_data["labels"]
                colors = ["#2563eb", "#dc2626", "#22c55e", "#f59e0b", "#8b5cf6",
                          "#ec4899", "#14b8a6", "#f97316", "#06b6d4", "#84cc16"]
                datasets = balance_data.get("datasets", [])
                # Detect if normalization is needed (same logic as program view)
                max_per = [max((abs(v) for v in ds.get("data", [])), default=0.0001) or 0.0001
                           for ds in datasets]
                g_max = max(max_per) if max_per else 1
                g_min = min(max_per) if max_per else 1
                needs_norm = g_max > 0 and g_min > 0 and (g_max / g_min) > 50
                if needs_norm:
                    log_max_per = [_math.log10(m + 1) for m in max_per]
                    log_g_max = max(log_max_per) if log_max_per else 1
                for i, ds in enumerate(datasets):
                    raw = ds.get("data", [])
                    token_label = ds.get("token", "?")
                    if needs_norm:
                        t_max = max_per[i] or 1
                        log_scale = log_max_per[i] / (log_g_max or 1)
                        ceiling = 20 + log_scale * 80
                        plot_data = [((v or 0) / t_max) * ceiling for v in raw]
                        lw = max(1, min(3, log_scale * 3))
                        token_label += " (skala)"
                    else:
                        plot_data = raw
                        lw = 1.2
                    ax.plot(range(len(labels)), plot_data,
                            label=token_label,
                            color=colors[i % len(colors)], linewidth=lw)
                # Show ~10 x-tick labels
                step = max(1, len(labels) // 10)
                ax.set_xticks(range(0, len(labels), step))
                ax.set_xticklabels([labels[j] for j in range(0, len(labels), step)],
                                   rotation=45, fontsize=7, ha="right")
                ax.set_title("Saldo w czasie", fontsize=11, fontweight="bold")
                if needs_norm:
                    ax.set_ylim(0, 105)
                    ax.set_ylabel("Skala relatywna (log)", fontsize=8)
                    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.0f}%"))
                ax.legend(fontsize=7, loc="upper left", ncol=3)
                ax.tick_params(axis="y", labelsize=8)
                ax.grid(True, alpha=0.3)
                fig.tight_layout()
                buf = BytesIO()
                fig.savefig(buf, format="png", dpi=150)
                plt.close(fig)
                buf.seek(0)
                doc.add_picture(buf, width=Cm(16))
                buf.close()

            if has_graph:
                import math
                nodes = graph.get("nodes", [])
                edges = graph.get("edges", [])
                risk_cm = {"critical": "#dc2626", "high": "#f97316", "medium": "#eab308", "low": "#64748b"}
                n_nodes = len(nodes)
                node_pos = {}
                for idx, node in enumerate(nodes):
                    angle = 2 * math.pi * idx / max(n_nodes, 1)
                    node_pos[node["data"]["id"]] = (math.cos(angle), math.sin(angle))

                fig, ax = plt.subplots(figsize=(7, 5))
                ax.set_aspect("equal")
                ax.axis("off")
                ax.set_title("Graf przepływu transakcji", fontsize=11, fontweight="bold")
                # Draw edges
                for edge in edges[:200]:
                    ed = edge.get("data", {})
                    s, t = ed.get("source", ""), ed.get("target", "")
                    if s in node_pos and t in node_pos:
                        x1, y1 = node_pos[s]
                        x2, y2 = node_pos[t]
                        ec = "#ef4444" if ed.get("risk") else "#c0c0c0"
                        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                                    arrowprops=dict(arrowstyle="->", color=ec, lw=0.6, alpha=0.5))
                # Draw nodes
                for node in nodes:
                    nd = node.get("data", {})
                    nid = nd.get("id", "")
                    if nid in node_pos:
                        x, y = node_pos[nid]
                        nc = risk_cm.get(nd.get("risk_level", "low"), "#64748b")
                        ax.plot(x, y, "o", color=nc, markersize=8, markeredgecolor="white", markeredgewidth=0.5)
                        ax.text(x, y - 0.08, nd.get("label", nid[:8]), ha="center", fontsize=5, color="#334155")
                fig.tight_layout()
                buf = BytesIO()
                fig.savefig(buf, format="png", dpi=150)
                plt.close(fig)
                buf.seek(0)
                doc.add_picture(buf, width=Cm(14))
                buf.close()
        except ImportError:
            pass  # matplotlib not available - skip charts in DOCX

    # ── 7. Numery telefonów ──
    phones = r.get("detected_phones", [])
    if phones:
        doc.add_heading("7. Zidentyfikowane numery telefonów", level=1)
        _add_desc("telefony")
        p_rows = [[p["number"], p.get("country_name", "?"), p.get("country_iso", "?"),
                    str(p.get("occurrences", 0))] for p in phones]
        add_data_table(["Numer", "Kraj", "ISO", "Wystąpienia"], p_rows)

    # ── 7. Kontrahenci ──
    cps = bs.get("counterparties", {})
    if cps:
        doc.add_heading("8. Kontrahenci wewnętrzni", level=1)
        _add_desc("kontrahenci")
        c_rows = []
        for uid, c in sorted(cps.items(), key=lambda x: x[1].get("tx_count", 0), reverse=True):
            c_rows.append([uid, str(c.get("tx_count", 0)),
                           f"{c.get('total_in', 0):.4f}", f"{c.get('total_out', 0):.4f}",
                           ", ".join(c.get("tokens", []))])
        add_data_table(["User ID", "TX", "Wpływy", "Wypływy", "Tokeny"], c_rows)

    # ── 8. Adresy zewnętrzne ──
    ext_src = fr.get("external_source_addresses", [])
    ext_dst = fr.get("external_dest_addresses", [])
    if ext_src or ext_dst:
        doc.add_heading("9. Adresy zewnętrzne", level=1)
        _add_desc("adresy")

        def _nk_docx(addr: str) -> str:
            a = addr.strip()
            return a.lower() if a.startswith("0x") or a.startswith("0X") else a

        addr_m: Dict[str, Dict[str, Any]] = {}
        for a in ext_src:
            key = _nk_docx(a["address"])
            if key in addr_m:
                addr_m[key]["dc"] += a["count"]
                addr_m[key]["dt"] += a["total"]
                addr_m[key]["tok"].update(a.get("tokens", []))
            else:
                addr_m[key] = {"disp": a["address"], "dc": a["count"], "dt": a["total"],
                               "wc": 0, "wt": 0.0, "tok": set(a.get("tokens", []))}
        for a in ext_dst:
            key = _nk_docx(a["address"])
            if key in addr_m:
                addr_m[key]["wc"] += a["count"]
                addr_m[key]["wt"] += a["total"]
                addr_m[key]["tok"].update(a.get("tokens", []))
            else:
                addr_m[key] = {"disp": a["address"], "dc": 0, "dt": 0.0,
                               "wc": a["count"], "wt": a["total"], "tok": set(a.get("tokens", []))}
        a_rows = []
        for _key, m in sorted(addr_m.items(), key=lambda x: x[1]["dt"] + x[1]["wt"], reverse=True):
            d = "DEP+WD" if m["dc"] > 0 and m["wc"] > 0 else ("DEP" if m["dc"] > 0 else "WD")
            a_rows.append([m["disp"], d, str(m["dc"]), f"{m['dt']:.4f}",
                           str(m["wc"]), f"{m['wt']:.4f}", ", ".join(sorted(m["tok"]))])
        add_data_table(["Adres", "Kier.", "Dep TX", "Dep suma", "Wyp TX", "Wyp suma", "Tokeny"], a_rows)

    # ── 9. Analiza czasowa ──
    d_ta = fr.get("temporal_analysis", {})
    if d_ta and d_ta.get("active_span_days"):
        doc.add_heading("10. Analiza czasowa", level=1)
        _add_desc("czasowa")
        add_kv_table([
            ("Okres aktywności", f"{d_ta.get('active_span_days', 0)} dni"),
            ("Aktywne dni", f"{d_ta.get('active_days', 0)} ({d_ta.get('activity_density', 0)}%)"),
            ("Szczytowa godzina", f"{d_ta.get('peak_hour', '?')}:00 ({d_ta.get('peak_hour_count', 0)} tx)"),
            ("Nocna aktywność", f"{d_ta.get('night_activity_count', 0)} ({d_ta.get('night_activity_ratio', 0)}%)"),
            ("Weekend/robocze", f"{d_ta.get('weekend_count', 0)}/{d_ta.get('weekday_count', 0)} ({d_ta.get('weekend_ratio', 0)}%)"),
        ])
        if d_ta.get("burst_days"):
            add_data_table(["Data", "TX"], [[b["date"], str(b["tx_count"])] for b in d_ta["burst_days"]])
        if d_ta.get("dormancy_periods"):
            add_data_table(["Od", "Do", "Dni"], [[d["from"], d["to"], str(d["days"])] for d in d_ta["dormancy_periods"]])

    # ── 10. Łańcuchy konwersji ──
    d_cc = fr.get("conversion_chains", {})
    if d_cc and d_cc.get("edges"):
        doc.add_heading("11. Łańcuchy konwersji tokenów", level=1)
        _add_desc("konwersje")
        doc.add_paragraph(f"Unikalne pary: {d_cc.get('unique_swap_pairs', 0)}")
        add_data_table(["Z tokenu", "Na token", "Wolumen"],
                       [[e["from"], e["to"], f"{e['volume']:.4f}"] for e in d_cc["edges"]])

    # ── 11. Structuring ──
    d_sd = fr.get("structuring_detection", {})
    if d_sd.get("alerts"):
        doc.add_heading("12. Wykrywanie structuringu", level=1)
        _add_desc("structuring")
        add_data_table(["Data", "Typ", "Próg", "TX", "Suma"],
                       [[a["date"], a["type"], str(a["threshold"]), str(a["count"]),
                         f"{a['daily_total']:.2f}"] for a in d_sd["alerts"]])

    # ── 12. Wash trading ──
    d_wt = fr.get("wash_trading", {})
    if d_wt.get("zero_net_markets"):
        doc.add_heading("13. Wash trading", level=1)
        _add_desc("wash")
        add_data_table(["Rynek", "Wol. brutto", "Poz. netto", "Net%", "Kupno", "Sprzedaż"],
                       [[m["market"], f"{m['gross_volume']:.4f}", f"{m['net_position']:.4f}",
                         f"{m['net_ratio']}%", f"{m['buys']:.4f}", f"{m['sells']:.4f}"]
                        for m in d_wt["zero_net_markets"]])

    # ── 13. Fiat ramp ──
    d_fa = fr.get("fiat_ramp_analysis", {})
    if d_fa and (d_fa.get("fiat_deposit_count", 0) > 0 or d_fa.get("fiat_withdrawal_count", 0) > 0):
        doc.add_heading("14. Analiza fiat on/off ramp", level=1)
        _add_desc("fiat")
        kv = [
            ("Wpłaty fiat", d_fa.get("fiat_deposit_count", 0)),
            ("Wypłaty fiat", d_fa.get("fiat_withdrawal_count", 0)),
            ("Łącznie IN", f"{d_fa.get('total_fiat_in', 0):.2f}"),
            ("Łącznie OUT", f"{d_fa.get('total_fiat_out', 0):.2f}"),
            ("Saldo netto", f"{d_fa.get('net_fiat_flow', 0):.2f}"),
        ]
        if d_fa.get("fiat_to_crypto_wd_hours") is not None:
            kv.append(("Fiat→crypto", f"{d_fa['fiat_to_crypto_wd_hours']:.1f} godz."))
        add_kv_table(kv)

    # ── 14. P2P ──
    d_p2p = fr.get("p2p_analysis", {})
    if d_p2p and d_p2p.get("total_count", 0) > 0:
        doc.add_heading("15. Analiza P2P", level=1)
        _add_desc("p2p")
        add_kv_table([
            ("Transakcje P2P", d_p2p["total_count"]),
            ("% aktywności", f"{d_p2p.get('total_pct', 0)}%"),
            ("Wolumen", f"{d_p2p.get('total_volume', 0):.2f}"),
            ("Kontrahenci", d_p2p.get("unique_counterparties", 0)),
        ])
        tops = d_p2p.get("top_counterparties", [])
        if tops:
            add_data_table(["ID", "TX", "Wolumen", "Tokeny"],
                           [[cp["id"], str(cp["count"]), f"{cp['volume']:.4f}",
                             ", ".join(cp.get("tokens", []))] for cp in tops])

    # ── 15. Velocity ──
    d_va = fr.get("velocity_analysis", {})
    if d_va and d_va.get("token_velocities"):
        doc.add_heading("16. Prędkość przepływu środków", level=1)
        _add_desc("velocity")
        add_kv_table([
            ("DEP/WD", f"{d_va.get('deposit_count', 0)}/{d_va.get('withdrawal_count', 0)}"),
            ("Ratio", d_va.get("dep_wd_ratio", 0)),
            ("Hot wallet", "TAK" if d_va.get("has_hot_wallet_behavior") else "NIE"),
        ])
        add_data_table(["Token", "Śr. godz.", "Min godz.", "Wpłaty", "Wypłaty"],
                       [[t["token"], str(t["avg_hold_hours"]), str(t["min_hold_hours"]),
                         str(t["deposit_count"]), str(t["withdrawal_count"])]
                        for t in d_va["token_velocities"]])

    # ── 16. Fee analysis ──
    d_fee = fr.get("fee_analysis", {})
    if d_fee and d_fee.get("fee_paying_tx_count", 0) > 0:
        doc.add_heading("17. Analiza opłat", level=1)
        _add_desc("fees")
        add_kv_table([
            ("TX z opłatami", d_fee["fee_paying_tx_count"]),
            ("Opłaty BNB", f"{d_fee.get('bnb_fee_count', 0)} ({d_fee.get('bnb_fee_ratio', 0)}%)"),
        ])
        fees_d = d_fee.get("total_fees_by_token", {})
        if fees_d:
            add_data_table(["Token", "Suma"],
                           [[tok, f"{val:.8f}"] for tok, val in fees_d.items()])

    # ── 17. Network analysis ──
    d_na = fr.get("network_analysis", {})
    if d_na and d_na.get("networks"):
        doc.add_heading("18. Analiza sieci blockchain", level=1)
        _add_desc("sieci")
        add_data_table(["Sieć", "Wpłaty", "Wypłaty", "TX", "Wol. wpłat", "Wol. wypłat"],
                       [[n["network"], str(n["deposits"]), str(n["withdrawals"]),
                         str(n["total_tx"]), f"{n['dep_volume']:.4f}", f"{n['wd_volume']:.4f}"]
                        for n in d_na["networks"]])

    # ── 18. Extended security ──
    d_es = fr.get("extended_security", {})
    if d_es:
        doc.add_heading("19. Rozszerzona analiza bezpieczeństwa", level=1)
        _add_desc("ext_security")
        add_kv_table([
            ("Kraje logowań", ", ".join(d_es.get("login_countries", []))),
            ("Podejrzane dni VPN", d_es.get("vpn_suspect_days", 0)),
            ("API Trading", "Tak" if d_es.get("api_trading_enabled") else "Nie"),
            ("Sub-konto", "Tak" if d_es.get("has_sub_account") else "Nie"),
        ])
        vpn_d = d_es.get("vpn_suspects", [])
        if vpn_d:
            add_data_table(["Data", "Kraje", "Ilość", "Loginy"],
                           [[v["date"], ", ".join(v.get("countries", [])),
                             str(v["country_count"]), str(v["login_count"])] for v in vpn_d])

    # ── 19. Transakcje ──
    txs = r.get("transactions", [])
    if txs:
        doc.add_heading("20. Transakcje", level=1)
        _add_desc("transakcje")
        doc.add_paragraph(f"Łącznie: {len(txs)}")
        tx_rows = []
        for tx in txs:
            ts = (tx.get("timestamp", "") or "")[:16]
            cp = tx.get("counterparty", "") or tx.get("to", "") or ""
            tx_rows.append([ts, tx.get("tx_type", ""), tx.get("token", ""),
                            f"{tx.get('amount', 0):.4f}", cp])
        add_data_table(["Data", "Typ", "Token", "Kwota", "Kontrahent/Do"], tx_rows)

    # ── Słownik typów transakcji ──
    used_types: set = set()
    for tx in r.get("transactions", []):
        tt = tx.get("tx_type", "")
        if tt:
            used_types.add(tt.lower())
        cat = tx.get("category", "")
        if cat:
            used_types.add(cat.lower())

    legend_items = [(tt, _TX_TYPE_DESCRIPTIONS[tt]) for tt in sorted(used_types) if tt in _TX_TYPE_DESCRIPTIONS]
    if legend_items:
        doc.add_heading("Słownik typów transakcji", level=1)
        add_data_table(["Typ", "Opis"], legend_items)

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph(f"Wygenerowano: {now} - AISTATE Crypto Analysis Module")
    p.runs[0].italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
