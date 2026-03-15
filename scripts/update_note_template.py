"""
Script to add Jinja2 chart image placeholders to the GSM note DOCX template.

Adds {%p if chart_X %} / {{ chart_X }} / {%p endif %} blocks after section headings.
Tables are NOT added here — they are inserted programmatically in note_generator.py.
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}

def qn(tag: str) -> str:
    prefix, local = tag.split(':')
    return f'{{{NSMAP[prefix]}}}{local}'


def get_body_index(body, element) -> int:
    for i, child in enumerate(body):
        if child is element:
            return i
    return -1


def find_section_last_element(body, heading_elem):
    """Find last body-level element in a section (before next Heading)."""
    started = False
    last = heading_elem
    for child in body:
        if child is heading_elem:
            started = True
            continue
        if started:
            if child.tag == qn('w:p'):
                pStyle = child.find(f'.//{qn("w:pStyle")}')
                if pStyle is not None and 'Heading' in pStyle.get(qn('w:val'), ''):
                    break
            last = child
    return last


def insert_chart_block(doc, body, after_elem, chart_var: str, caption: str):
    """Insert a conditional chart image block using {%p if %} syntax.

    This creates 3 paragraphs:
      {%p if chart_var %}
      <caption bold centered>
      {{ chart_var }}   (centered — InlineImage placeholder)
      {%p endif %}

    Using {%p ...} means docxtpl will remove the entire paragraph if condition is false.
    """
    # Paragraph 1: {%p if chart_var %}
    p1 = doc.add_paragraph()
    body.remove(p1._element)
    idx = get_body_index(body, after_elem)
    body.insert(idx + 1, p1._element)
    r1 = p1.add_run('{%p if ' + chart_var + ' %}')
    r1.font.size = Pt(1)
    r1.font.color.rgb = RGBColor(255, 255, 255)

    # Paragraph 2: Caption
    p2 = doc.add_paragraph()
    body.remove(p2._element)
    idx = get_body_index(body, p1._element)
    body.insert(idx + 1, p2._element)
    r2 = p2.add_run(caption)
    r2.bold = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0, 0, 0)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Paragraph 3: Image placeholder {{ chart_var }}
    p3 = doc.add_paragraph()
    body.remove(p3._element)
    idx = get_body_index(body, p2._element)
    body.insert(idx + 1, p3._element)
    r3 = p3.add_run('{{ ' + chart_var + ' }}')
    r3.font.size = Pt(10)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Paragraph 4: {%p endif %}
    p4 = doc.add_paragraph()
    body.remove(p4._element)
    idx = get_body_index(body, p3._element)
    body.insert(idx + 1, p4._element)
    r4 = p4.add_run('{%p endif %}')
    r4.font.size = Pt(1)
    r4.font.color.rgb = RGBColor(255, 255, 255)

    return p4._element  # return last element for chaining


def update_template(input_path: Path, output_path: Path):
    """Add chart placeholders to the DOCX template."""
    doc = Document(str(input_path))
    body = doc.element.body

    # Find section headings
    sections = {}
    for p in doc.paragraphs:
        text = p.text.strip()
        if p.style.name.startswith('Heading'):
            if text.startswith('4.'):
                sections['4'] = p._element
            elif text.startswith('5.'):
                sections['5'] = p._element
            elif text.startswith('6.'):
                sections['6'] = p._element
            elif text.startswith('7.'):
                sections['7'] = p._element

    print(f"Found section headings: {list(sections.keys())}")

    # Section 4: Activity charts (after last paragraph of section 4)
    if '4' in sections:
        last = find_section_last_element(body, sections['4'])
        print("  Adding charts after section 4: activity, night, weekend")
        last = insert_chart_block(doc, body, last, 'chart_activity',
            'Wykres: Rozkład aktywności telekomunikacyjnej')
        last = insert_chart_block(doc, body, last, 'chart_night_activity',
            'Wykres: Aktywność nocna')
        last = insert_chart_block(doc, body, last, 'chart_weekend_activity',
            'Wykres: Aktywność weekendowa')

    # Section 5: Contacts chart
    if '5' in sections:
        last = find_section_last_element(body, sections['5'])
        print("  Adding chart after section 5: top_contacts")
        insert_chart_block(doc, body, last, 'chart_top_contacts',
            'Wykres: Najczęstsze kontakty')

    # Section 7: BTS map
    if '7' in sections:
        last = find_section_last_element(body, sections['7'])
        print("  Adding chart after section 7: map_bts")
        insert_chart_block(doc, body, last, 'chart_map_bts',
            'Mapa: Lokalizacje stacji BTS')

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"\nSaved: {output_path}")

    # Verify template variables
    from docxtpl import DocxTemplate
    tpl = DocxTemplate(str(output_path))
    try:
        vars_ = tpl.get_undeclared_template_variables()
        print(f"Variables ({len(vars_)}): {sorted(vars_)}")
    except Exception as e:
        print(f"Variable scan warning (may work at render time): {e}")


if __name__ == '__main__':
    root = Path(__file__).resolve().parent.parent
    backup = root / 'templates' / 'gsm_note_template.docx.bak'
    output = root / 'templates' / 'gsm_note_template.docx'

    if not backup.exists():
        print(f"Backup not found: {backup}")
        sys.exit(1)

    # Always start from clean backup
    shutil.copy2(backup, output)
    print(f"Restored from backup: {backup}")

    update_template(output, output)
